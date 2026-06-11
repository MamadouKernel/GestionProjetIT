from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ACCENT = RGBColor(30, 64, 175)
TEXT = RGBColor(31, 41, 55)
MUTED = RGBColor(75, 85, 99)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_borders(cell, color="D1D5DB"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        edge_el = borders.find(qn(f"w:{edge}"))
        if edge_el is None:
            edge_el = OxmlElement(f"w:{edge}")
            borders.append(edge_el)
        edge_el.set(qn("w:val"), "single")
        edge_el.set(qn("w:sz"), "6")
        edge_el.set(qn("w:space"), "0")
        edge_el.set(qn("w:color"), color)


def apply_base_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = TEXT
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    if "TemplateHeading1" not in doc.styles:
        style = doc.styles.add_style("TemplateHeading1", WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = doc.styles["Heading 1"]
        style.font.name = "Aptos Display"
        style.font.size = Pt(16)
        style.font.bold = True
        style.font.color.rgb = ACCENT
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(8)

    if "TemplateHeading2" not in doc.styles:
        style = doc.styles.add_style("TemplateHeading2", WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = doc.styles["Heading 2"]
        style.font.name = "Aptos"
        style.font.size = Pt(12)
        style.font.bold = True
        style.font.color.rgb = TEXT
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(4)

    if "Label" not in doc.styles:
        style = doc.styles.add_style("Label", WD_STYLE_TYPE.CHARACTER)
        style.font.bold = True
        style.font.color.rgb = TEXT

    if "Placeholder" not in doc.styles:
        style = doc.styles.add_style("Placeholder", WD_STYLE_TYPE.CHARACTER)
        style.font.italic = True
        style.font.color.rgb = MUTED


def add_cover(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("MODELE DE CAHIER DES CHARGES")
    run.bold = True
    run.font.name = "Aptos Display"
    run.font.size = Pt(22)
    run.font.color.rgb = ACCENT

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Projet IT")
    run.font.name = "Aptos"
    run.font.size = Pt(16)
    run.font.color.rgb = TEXT

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "Template de saisie a joindre a la demande de projet dans l'application de gestion."
    )
    run.font.size = Pt(11)
    run.font.color.rgb = MUTED

    doc.add_paragraph("")

    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    rows = [
        ("Intitule du projet", "[A completer]"),
        ("Direction demandeuse", "[A completer]"),
        ("Demandeur", "[Nom et prenoms]"),
        ("Directeur metier", "[Nom et prenoms]"),
        ("Autre sponsor (CODIR)", "[Optionnel]"),
        ("Priorite calculee", "[P1 a P6 selon Urgence + Criticite]"),
        ("Date souhaitee de mise en oeuvre", "[JJ/MM/AAAA]"),
        ("Version du document", "v1.0"),
    ]
    for label, value in rows:
        row = table.add_row().cells
        row[0].width = Cm(6.2)
        row[1].width = Cm(9.8)
        row[0].paragraphs[0].add_run(label).bold = True
        row[1].paragraphs[0].add_run(value).italic = True
        for cell in row:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_borders(cell)
        set_cell_shading(row[0], "EFF6FF")

    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.LEFT
    note.add_run("Instruction : ").bold = True
    note.add_run(
        "remplir les rubriques obligatoires, joindre les annexes utiles, puis exporter ou enregistrer le document avant depot dans l'application."
    )


def add_section_title(doc, title):
    doc.add_paragraph(title, style="TemplateHeading1")


def add_prompt_paragraph(doc, label, placeholder):
    p = doc.add_paragraph()
    p.add_run(f"{label} : ", style="Label")
    p.add_run(placeholder, style="Placeholder")


def add_bullet_list(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_multiline_block(doc, title, lines=4):
    doc.add_paragraph(title, style="TemplateHeading2")
    for _ in range(lines):
        p = doc.add_paragraph()
        p.add_run("....................................................................................................................")
        p.runs[0].font.color.rgb = MUTED


def build_document(output_path: Path):
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

    apply_base_styles(doc)
    add_cover(doc)

    doc.add_page_break()

    add_section_title(doc, "1. Resume du besoin")
    add_prompt_paragraph(doc, "Problematique metier", "[Decrire clairement le besoin initial, son origine et les irritants constates]")
    add_prompt_paragraph(doc, "Contexte", "[Presenter le contexte, les enjeux et les dependencies connues]")
    add_prompt_paragraph(doc, "Objectifs", "[Lister les objectifs mesurables et attendus du projet]")
    add_prompt_paragraph(doc, "Avantages attendus", "[Expliquer les gains metier, organisationnels, qualitatifs ou reglementaires]")

    add_section_title(doc, "2. Perimetre et parties prenantes")
    add_prompt_paragraph(doc, "Perimetre", "[Preciser les processus impactes, les activites couvertes et les populations concernees]")
    add_prompt_paragraph(doc, "Hors perimetre", "[Lister explicitement ce qui n'est pas inclus dans la demande]")
    add_prompt_paragraph(doc, "Directions et equipes impliquees", "[Identifier les directions contributrices et les principaux interlocuteurs]")
    add_prompt_paragraph(doc, "Autre sponsor (CODIR)", "[Optionnel - a renseigner si un sponsor complementaire soutient la demande]")

    add_section_title(doc, "3. Priorite et contraintes")
    add_prompt_paragraph(doc, "Urgence", "[Basse / Moyenne / Haute - justifier le niveau choisi]")
    add_prompt_paragraph(doc, "Criticite", "[Faible / Moyenne / Elevee / Critique - preciser l'impact en cas de non realisation]")
    add_prompt_paragraph(doc, "Priorite calculee", "[La priorite P1 a P6 est calculee automatiquement dans l'application a partir de l'urgence et de la criticite]")
    add_prompt_paragraph(doc, "Contraintes", "[Reglementaires, techniques, organisationnelles, calendaires, ressources]")
    add_prompt_paragraph(doc, "Date souhaitee de mise en oeuvre", "[JJ/MM/AAAA]")

    doc.add_section(WD_SECTION.NEW_PAGE)

    add_section_title(doc, "4. Livrables et criteres de succes")
    add_bullet_list(
        doc,
        [
            "Livrable 1 : [A completer]",
            "Livrable 2 : [A completer]",
            "Livrable 3 : [A completer]",
        ],
    )
    add_prompt_paragraph(doc, "Criteres de succes", "[Definir les conditions qui permettront de considerer le besoin comme satisfait]")
    add_prompt_paragraph(doc, "Indicateurs de suivi", "[Proposer, si possible, des indicateurs simples de mesure du resultat]")

    add_section_title(doc, "5. Risques, hypotheses et dependances")
    add_prompt_paragraph(doc, "Risques initiaux", "[Lister les risques principaux et les mesures de mitigation envisagees]")
    add_prompt_paragraph(doc, "Hypotheses", "[Lister les hypotheses structurantes pour la demande]")
    add_prompt_paragraph(doc, "Dependances", "[Projets, applications, fournisseurs, arbitrages ou validations externes]")

    add_section_title(doc, "6. Annexes et validation")
    add_prompt_paragraph(doc, "Annexes jointes", "[Liste des documents annexes a transmettre avec la demande]")
    add_prompt_paragraph(doc, "Commentaires complementaires", "[Informations utiles non couvertes dans les sections precedentes]")

    approvals = doc.add_table(rows=1, cols=3)
    approvals.alignment = WD_TABLE_ALIGNMENT.CENTER
    approvals.autofit = False
    headers = approvals.rows[0].cells
    headers[0].text = "Role"
    headers[1].text = "Nom"
    headers[2].text = "Visa / date"
    for cell in headers:
        cell.paragraphs[0].runs[0].bold = True
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_borders(cell)
        set_cell_shading(cell, "DBEAFE")

    for role in ("Demandeur", "Directeur metier", "DSI"):
        row = approvals.add_row().cells
        row[0].text = role
        row[1].text = "[A completer]"
        row[2].text = "[A completer]"
        for cell in row:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_borders(cell)

    doc.add_paragraph("")
    footer_note = doc.add_paragraph()
    footer_note.alignment = WD_ALIGN_PARAGRAPH.LEFT
    footer_note.add_run("Rappel : ").bold = True
    footer_note.add_run(
        "les rubriques Description, Contexte, Objectifs, Avantages attendus et Cahier des charges sont obligatoires dans l'application."
    )

    doc.save(output_path)


if __name__ == "__main__":
    base_dir = Path(__file__).resolve().parents[1]
    output_dir = base_dir / "wwwroot" / "templates"
    output_dir.mkdir(parents=True, exist_ok=True)
    output_path = output_dir / "Modele_Cahier_Charges_Projet_IT.docx"
    build_document(output_path)
    print(output_path)
