# -*- coding: utf-8 -*-
"""
Generation du document Word : Workflow complet par role
Gestion Projets IT - Cote d'Ivoire Terminal
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import date
import copy

# Palette CIT
NAVY  = RGBColor(0x16, 0x23, 0x47)
BLUE  = RGBColor(0x1A, 0x72, 0xB8)
GOLD  = RGBColor(0xC9, 0xA2, 0x27)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT = RGBColor(0xEE, 0xF4, 0xFB)
GRAY  = RGBColor(0xF5, 0xF5, 0xF5)
GREEN = RGBColor(0x10, 0xB9, 0x81)
RED   = RGBColor(0xEF, 0x44, 0x44)

def hex_to_rgb_str(r, g, b):
    return f"{r:02X}{g:02X}{b:02X}"

NAVY_HEX  = "162347"
BLUE_HEX  = "1A72B8"
GOLD_HEX  = "C9A227"
WHITE_HEX = "FFFFFF"
LIGHT_HEX = "EEF4FB"

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_borders(cell, color="CCCCCC"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), color)
        tcBorders.append(border)
    tcPr.append(tcBorders)

def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    if level == 1:
        run.font.size = Pt(18)
        run.font.bold = True
        run.font.color.rgb = NAVY
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(6)
        # Ligne sous le titre
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '8')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), GOLD_HEX)
        pBdr.append(bottom)
        pPr.append(pBdr)
    elif level == 2:
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = BLUE
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
    elif level == 3:
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = NAVY
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)
    return p

def add_para(doc, text, indent=False, bullet=False, color=None, bold=False, italic=False):
    p = doc.add_paragraph()
    if bullet:
        p.style = doc.styles['List Bullet']
    if indent:
        p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    p.paragraph_format.space_after = Pt(3)
    return p

def add_step_table(doc, steps):
    """Tableau de workflow avec etapes numerotees."""
    table = doc.add_table(rows=len(steps), cols=3)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    col_widths = [Cm(1.5), Cm(5.5), Cm(10)]

    for i, (num, actor, description) in enumerate(steps):
        row = table.rows[i]
        row.height = Cm(1.0)

        # Cellule numero
        c0 = row.cells[0]
        c0.width = col_widths[0]
        set_cell_bg(c0, BLUE_HEX if i % 2 == 0 else NAVY_HEX)
        set_cell_borders(c0, BLUE_HEX)
        p0 = c0.paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r0 = p0.add_run(str(num))
        r0.font.bold = True
        r0.font.color.rgb = WHITE
        r0.font.size = Pt(13)
        c0.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # Cellule acteur
        c1 = row.cells[1]
        c1.width = col_widths[1]
        set_cell_bg(c1, LIGHT_HEX)
        set_cell_borders(c1, "CCCCCC")
        p1 = c1.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p1.add_run(actor)
        r1.font.bold = True
        r1.font.color.rgb = NAVY
        r1.font.size = Pt(10)
        c1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # Cellule description
        c2 = row.cells[2]
        c2.width = col_widths[2]
        set_cell_bg(c2, WHITE_HEX)
        set_cell_borders(c2, "CCCCCC")
        p2 = c2.paragraphs[0]
        r2 = p2.add_run(description)
        r2.font.size = Pt(10)
        r2.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
        c2.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    doc.add_paragraph()
    return table

def add_decision_box(doc, condition, yes_text, no_text=None):
    """Boite de decision oui/non."""
    table = doc.add_table(rows=1, cols=3 if no_text else 2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Condition
    c0 = table.rows[0].cells[0]
    c0.width = Cm(5)
    set_cell_bg(c0, GOLD_HEX)
    set_cell_borders(c0, GOLD_HEX)
    p = c0.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(condition)
    r.font.bold = True
    r.font.color.rgb = WHITE
    r.font.size = Pt(10)
    c0.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Oui
    c1 = table.rows[0].cells[1]
    c1.width = Cm(6)
    set_cell_bg(c1, "D4EDDA")
    set_cell_borders(c1, "28A745")
    p1 = c1.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p1.add_run("OUI  =>  " + yes_text)
    r1.font.bold = True
    r1.font.color.rgb = RGBColor(0x15, 0x52, 0x24)
    r1.font.size = Pt(10)
    c1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    if no_text:
        c2 = table.rows[0].cells[2]
        c2.width = Cm(6)
        set_cell_bg(c2, "F8D7DA")
        set_cell_borders(c2, "DC3545")
        p2 = c2.paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run("NON  =>  " + no_text)
        r2.font.bold = True
        r2.font.color.rgb = RGBColor(0x72, 0x1C, 0x24)
        r2.font.size = Pt(10)
        c2.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    doc.add_paragraph()

def add_info_box(doc, title, content, color_hex=LIGHT_HEX, title_color=NAVY_HEX):
    table = doc.add_table(rows=2, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    c_title = table.rows[0].cells[0]
    set_cell_bg(c_title, title_color)
    set_cell_borders(c_title, title_color)
    p = c_title.paragraphs[0]
    r = p.add_run(title)
    r.font.bold = True
    r.font.color.rgb = WHITE
    r.font.size = Pt(10)

    c_body = table.rows[1].cells[0]
    set_cell_bg(c_body, color_hex)
    set_cell_borders(c_body, "CCCCCC")
    p2 = c_body.paragraphs[0]
    r2 = p2.add_run(content)
    r2.font.size = Pt(10)
    r2.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    p2.paragraph_format.space_before = Pt(4)
    p2.paragraph_format.space_after = Pt(4)
    c_body.paragraph_format = p2.paragraph_format

    doc.add_paragraph()

def add_page_break(doc):
    doc.add_page_break()

def add_role_header(doc, role, icon_text, description, permissions):
    """En-tete visuel pour chaque role."""
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    c_icon = table.rows[0].cells[0]
    c_icon.width = Cm(3)
    set_cell_bg(c_icon, NAVY_HEX)
    set_cell_borders(c_icon, NAVY_HEX)
    p = c_icon.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(icon_text)
    r.font.size = Pt(24)
    r.font.color.rgb = GOLD
    c_icon.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    c_info = table.rows[0].cells[1]
    c_info.width = Cm(14)
    set_cell_bg(c_info, LIGHT_HEX)
    set_cell_borders(c_info, BLUE_HEX)
    p2 = c_info.paragraphs[0]
    r2 = p2.add_run(role)
    r2.font.bold = True
    r2.font.size = Pt(16)
    r2.font.color.rgb = NAVY
    p3 = c_info.add_paragraph()
    r3 = p3.add_run(description)
    r3.font.size = Pt(10)
    r3.font.italic = True
    r3.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    p4 = c_info.add_paragraph()
    r4 = p4.add_run("Permissions : " + permissions)
    r4.font.size = Pt(9)
    r4.font.color.rgb = BLUE

    doc.add_paragraph()

# ─────────────────────────────────────────────────────────────────────────────

def build_document():
    doc = Document()

    # Marges
    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # Style par defaut
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # ── PAGE DE COUVERTURE ────────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(60)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("GESTION PROJETS IT")
    r.font.size = Pt(28)
    r.font.bold = True
    r.font.color.rgb = NAVY

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("Cote d'Ivoire Terminal")
    r2.font.size = Pt(18)
    r2.font.color.rgb = GOLD
    r2.font.bold = True

    doc.add_paragraph()

    table_cover = doc.add_table(rows=1, cols=1)
    table_cover.alignment = WD_TABLE_ALIGNMENT.CENTER
    c = table_cover.rows[0].cells[0]
    c.width = Cm(14)
    set_cell_bg(c, NAVY_HEX)
    set_cell_borders(c, NAVY_HEX)
    p3 = c.paragraphs[0]
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run("WORKFLOW COMPLET PAR ROLE")
    r3.font.size = Pt(20)
    r3.font.bold = True
    r3.font.color.rgb = WHITE
    p4 = c.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r4 = p4.add_run("Guide d'utilisation detaille — A a Z")
    r4.font.size = Pt(12)
    r4.font.color.rgb = GOLD
    r4.font.italic = True

    for _ in range(3):
        doc.add_paragraph()

    info_table = doc.add_table(rows=4, cols=2)
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    infos = [
        ("Version", "1.0"),
        ("Date", date.today().strftime("%d/%m/%Y")),
        ("Projet", "Gestion Projets IT"),
        ("Statut", "Document de reference"),
    ]
    for i, (label, val) in enumerate(infos):
        cl = info_table.rows[i].cells[0]
        cv = info_table.rows[i].cells[1]
        set_cell_bg(cl, BLUE_HEX)
        set_cell_bg(cv, LIGHT_HEX)
        set_cell_borders(cl, BLUE_HEX)
        set_cell_borders(cv, "CCCCCC")
        pl = cl.paragraphs[0]
        rl = pl.add_run(label)
        rl.font.bold = True
        rl.font.color.rgb = WHITE
        rl.font.size = Pt(11)
        pv = cv.paragraphs[0]
        rv = pv.add_run(val)
        rv.font.color.rgb = NAVY
        rv.font.size = Pt(11)

    add_page_break(doc)

    # ── SOMMAIRE ──────────────────────────────────────────────────────────────
    add_heading(doc, "Sommaire")
    roles_toc = [
        ("1.", "Vue d'ensemble de l'application"),
        ("2.", "Role : Demandeur"),
        ("3.", "Role : Directeur Metier (DM / Sponsor)"),
        ("4.", "Role : DSI"),
        ("5.", "Role : Responsable Solutions IT (Chef de Projet)"),
        ("6.", "Role : AdminIT"),
        ("7.", "Role : Chef de Projet"),
        ("8.", "Matrice des droits par role"),
        ("9.", "Workflow global — Vue synthetique"),
    ]
    for num, title in roles_toc:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        r_num = p.add_run(f"{num}  ")
        r_num.font.bold = True
        r_num.font.color.rgb = BLUE
        r_title = p.add_run(title)
        r_title.font.color.rgb = NAVY

    add_page_break(doc)

    # ── 1. VUE D'ENSEMBLE ────────────────────────────────────────────────────
    add_heading(doc, "1. Vue d'ensemble de l'application")
    add_para(doc, "L'application Gestion Projets IT de Cote d'Ivoire Terminal (CIT) permet de gerer l'integralite du cycle de vie des projets informatiques, depuis la soumission d'une demande jusqu'a la cloture formelle du projet.", bold=False)
    doc.add_paragraph()

    add_heading(doc, "6 roles utilisateur", level=2)
    roles_overview = [
        ("Demandeur",               "Soumet des demandes de projet", "Demande uniquement"),
        ("Directeur Metier (DM)",   "Valide/rejette les demandes de sa direction", "Demandes + supervision"),
        ("DSI",                     "Decision finale + creation projet + portefeuille complet", "Acces complet"),
        ("Responsable Solutions IT","Gere les projets (toutes phases)", "Projets complets"),
        ("Chef de Projet",          "Pilote son projet affecte", "Son projet uniquement"),
        ("AdminIT",                 "Administration de l'application", "Administration"),
    ]
    t = doc.add_table(rows=len(roles_overview)+1, cols=3)
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Role", "Responsabilite principale", "Perimetre"]
    for ci, h in enumerate(headers):
        c = t.rows[0].cells[ci]
        set_cell_bg(c, NAVY_HEX)
        set_cell_borders(c, NAVY_HEX)
        p = c.paragraphs[0]
        r = p.add_run(h)
        r.font.bold = True
        r.font.color.rgb = WHITE
        r.font.size = Pt(11)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i, (role, resp, perm) in enumerate(roles_overview):
        row = t.rows[i+1]
        vals = [role, resp, perm]
        bg = LIGHT_HEX if i % 2 == 0 else WHITE_HEX
        for ci, val in enumerate(vals):
            c = row.cells[ci]
            set_cell_bg(c, bg)
            set_cell_borders(c, "CCCCCC")
            p = c.paragraphs[0]
            r = p.add_run(val)
            r.font.size = Pt(10)
            if ci == 0:
                r.font.bold = True
                r.font.color.rgb = BLUE
    doc.add_paragraph()

    add_heading(doc, "5 phases du cycle de vie projet", level=2)
    phases = [
        ("Phase 1", "Analyse & Clarification",      "Note de cadrage, compte-rendus"),
        ("Phase 2", "Planification & Validation",   "Planning, budget, charte projet signee"),
        ("Phase 3", "Execution & Suivi",             "Taches, risques, livrables, bilan hebdo"),
        ("Phase 4", "UAT & MEP",                     "Tests utilisateurs, mise en production"),
        ("Phase 5", "Cloture & Lecons apprises",    "Rapport final, capitalisation"),
    ]
    t2 = doc.add_table(rows=len(phases)+1, cols=3)
    t2.style = 'Table Grid'
    t2.alignment = WD_TABLE_ALIGNMENT.CENTER
    for ci, h in enumerate(["Phase", "Nom", "Livrables cles"]):
        c = t2.rows[0].cells[ci]
        set_cell_bg(c, BLUE_HEX)
        set_cell_borders(c, BLUE_HEX)
        p = c.paragraphs[0]
        r = p.add_run(h)
        r.font.bold = True
        r.font.color.rgb = WHITE
        r.font.size = Pt(11)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i, (ph, name, liv) in enumerate(phases):
        row = t2.rows[i+1]
        bg = LIGHT_HEX if i % 2 == 0 else WHITE_HEX
        for ci, val in enumerate([ph, name, liv]):
            c = row.cells[ci]
            set_cell_bg(c, bg)
            set_cell_borders(c, "CCCCCC")
            p = c.paragraphs[0]
            r = p.add_run(val)
            r.font.size = Pt(10)
            if ci == 0:
                r.font.bold = True
                r.font.color.rgb = NAVY
    doc.add_paragraph()

    add_page_break(doc)

    # ─────────────────────────────────────────────────────────────────────────
    # ROLE : DEMANDEUR
    # ─────────────────────────────────────────────────────────────────────────
    add_heading(doc, "2. Role : Demandeur")
    add_role_header(doc, "DEMANDEUR",
        "[D]",
        "Collaborateur CIT souhaitant initier un projet informatique.",
        "Soumettre des demandes, suivre leur avancement, valider la cloture de son projet.")

    add_heading(doc, "A. Creer un compte (si nouveau)", level=2)
    add_step_table(doc, [
        (1, "Demandeur", "Acceder a la page de connexion de l'application"),
        (2, "Demandeur", "Cliquer sur 'Demander un acces' ou 'S'inscrire'"),
        (3, "Demandeur", "Remplir le formulaire : nom, prenom, email, direction, service, role souhaite"),
        (4, "Systeme",   "Email automatique envoye au Directeur Metier (DM) de la direction"),
        (5, "DM",        "Valide la demande de compte => DSI cree le compte"),
        (6, "Systeme",   "Email envoye au Demandeur avec son matricule et un lien d'activation securise"),
        (7, "Demandeur", "Definir son mot de passe via le lien d'activation puis se connecter"),
    ])

    add_heading(doc, "B. Soumettre une demande de projet", level=2)
    add_step_table(doc, [
        (1, "Demandeur", "Se connecter avec son matricule et mot de passe"),
        (2, "Demandeur", "Cliquer sur 'Nouvelle demande' dans le menu"),
        (3, "Demandeur", "Remplir le formulaire :\n- Titre du projet\n- Description et objectifs\n- Perimetre et benefices attendus\n- Priorite (Haute / Moyenne / Basse)\n- Budget estime\n- Direction (pre-remplie automatiquement)"),
        (4, "Demandeur", "Joindre des documents si necessaire"),
        (5, "Demandeur", "Cliquer sur 'Soumettre'"),
        (6, "Systeme",   "Statut : 'Soumise' — Email envoye au Directeur Metier"),
    ])
    add_decision_box(doc, "DM valide ?", "=> Demande envoyee a la DSI", "=> Demande rejetee / renvoyee")

    add_heading(doc, "C. Corriger une demande renvoyee", level=2)
    add_step_table(doc, [
        (1, "Demandeur", "Recevoir un email de renvoi avec le commentaire du DM"),
        (2, "Demandeur", "Se connecter et aller dans 'Mes demandes'"),
        (3, "Demandeur", "Ouvrir la demande renvoyee (statut : 'Renvoye demandeur')"),
        (4, "Demandeur", "Apporter les corrections demandees"),
        (5, "Demandeur", "Re-soumettre la demande"),
        (6, "Systeme",   "Email de notification envoye au DM"),
    ])

    add_heading(doc, "D. Suivre l'avancement de sa demande", level=2)
    add_step_table(doc, [
        (1, "Demandeur", "Aller dans 'Mes demandes'"),
        (2, "Demandeur", "Consulter le statut : Soumise / Validee DM / En attente DSI / Validee / Rejetee"),
        (3, "Demandeur", "Si le projet est cree, il apparait dans 'Mes projets' en lecture seule"),
    ])

    add_heading(doc, "E. Valider la cloture du projet", level=2)
    add_step_table(doc, [
        (1, "Systeme",   "Notification : le Chef de Projet a soumis une demande de cloture"),
        (2, "Demandeur", "Se connecter et aller dans 'Mes projets'"),
        (3, "Demandeur", "Ouvrir le projet en phase Cloture"),
        (4, "Demandeur", "Consulter le rapport de cloture"),
        (5, "Demandeur", "Valider ou rejeter la demande de cloture"),
    ])

    add_info_box(doc,
        "Acces Demandeur — Synthese",
        "Le Demandeur peut uniquement : creer/corriger ses demandes, suivre leur statut, consulter son projet en lecture seule, et valider la cloture.\nIl ne peut pas acceder aux projets d'autres directions ni modifier les donnees d'un projet.",
        LIGHT_HEX, BLUE_HEX)

    add_page_break(doc)

    # ─────────────────────────────────────────────────────────────────────────
    # ROLE : DIRECTEUR METIER
    # ─────────────────────────────────────────────────────────────────────────
    add_heading(doc, "3. Role : Directeur Metier (DM / Sponsor)")
    add_role_header(doc, "DIRECTEUR METIER (DM / SPONSOR)",
        "[DM]",
        "Responsable de direction. Valide les demandes de sa direction et supervise ses projets.",
        "Valider/rejeter les demandes, superviser les projets de sa direction, signer la charte, valider la recette.")

    add_heading(doc, "A. Valider ou rejeter une demande de projet", level=2)
    add_step_table(doc, [
        (1, "Systeme",  "Email recu : 'Nouvelle demande de projet — [Titre]'"),
        (2, "DM",       "Se connecter et aller dans 'Validation DM'"),
        (3, "DM",       "Ouvrir la demande et consulter tous les details"),
        (4, "DM",       "Choisir une action :"),
    ])
    add_decision_box(doc, "Decision DM ?",
        "VALIDER => Statut 'Validee DM', email DSI",
        "REJETER => Saisir motif, email Demandeur")
    add_step_table(doc, [
        (5, "DM",      "Option : Renvoyer au Demandeur avec commentaire pour correction"),
    ])

    add_heading(doc, "B. Superviser les projets de sa direction", level=2)
    add_step_table(doc, [
        (1, "DM", "Aller dans 'Mes projets' ou le 'Portefeuille'"),
        (2, "DM", "Voir tous les projets de sa direction (lecture seule si pas Sponsor)"),
        (3, "DM", "Consulter : statut, phase, avancement, risques, livrables"),
        (4, "DM", "En tant que Sponsor : possibilite d'agir sur les validations"),
    ])

    add_heading(doc, "C. Signer la charte projet (role Sponsor)", level=2)
    add_step_table(doc, [
        (1, "Systeme", "Notification : charte projet generee, signature requise"),
        (2, "DM",      "Se connecter et ouvrir le projet en phase Planification"),
        (3, "DM",      "Cliquer sur 'Signer la charte'"),
        (4, "DM",      "Apposer sa signature electronique (canvas)"),
        (5, "Systeme", "Signature enregistree — notification Chef de Projet"),
        (6, "Systeme", "Si toutes les signatures presentes => CharteValidee = true"),
    ])

    add_heading(doc, "D. Valider la recette (UAT) — GO/NO-GO MEP", level=2)
    add_step_table(doc, [
        (1, "Systeme", "Projet en phase UAT & MEP"),
        (2, "DM",      "Se connecter et ouvrir le projet, onglet 'UAT'"),
        (3, "DM",      "Consulter les resultats des tests et les anomalies"),
        (4, "DM",      "Cliquer sur 'Valider Recette' si satisfait"),
        (5, "Systeme", "RecetteValidee = true — Chef de Projet peut poursuivre"),
    ])
    add_decision_box(doc, "Recette OK ?",
        "VALIDER => passage possible en cloture",
        "REFUSER => retour en execution pour corrections")

    add_heading(doc, "E. Valider la cloture du projet", level=2)
    add_step_table(doc, [
        (1, "Systeme", "Notification : demande de cloture soumise par le Chef de Projet"),
        (2, "DM",      "Ouvrir le projet, onglet 'Cloture'"),
        (3, "DM",      "Lire le rapport de cloture et les lecons apprises"),
        (4, "DM",      "Valider la cloture"),
        (5, "Systeme", "Projet passe en statut 'Cloture'"),
    ])

    add_page_break(doc)

    # ─────────────────────────────────────────────────────────────────────────
    # ROLE : DSI
    # ─────────────────────────────────────────────────────────────────────────
    add_heading(doc, "4. Role : DSI")
    add_role_header(doc, "DSI — DIRECTION DES SYSTEMES D'INFORMATION",
        "[DSI]",
        "Responsable SI. Valide les demandes DM, cree les projets, gere le portefeuille complet.",
        "Acces complet : toutes les demandes, tous les projets, toutes les phases.")

    add_heading(doc, "A. Valider ou rejeter une demande (apres DM)", level=2)
    add_step_table(doc, [
        (1, "Systeme", "Email recu : 'Demande validee DM — en attente DSI — [Titre]'"),
        (2, "DSI",     "Aller dans 'Validation DSI'"),
        (3, "DSI",     "Consulter la demande, le metier, la faisabilite"),
        (4, "DSI",     "Choisir une action :"),
    ])
    add_decision_box(doc, "Decision DSI ?",
        "VALIDER => Projet cree automatiquement",
        "REJETER => Motif obligatoire, email DM + Demandeur")
    add_step_table(doc, [
        (5, "Systeme", "Si Validee : Code projet genere (ex: PRJ-2026-001)\nEmail envoye au DM et au Demandeur"),
    ])

    add_heading(doc, "B. Affecter un Chef de Projet", level=2)
    add_step_table(doc, [
        (1, "DSI", "Ouvrir le projet nouvellement cree"),
        (2, "DSI", "Aller dans la synthese du projet"),
        (3, "DSI", "Selectionner un Responsable Solutions IT dans la liste"),
        (4, "DSI", "Enregistrer — le Chef de Projet peut desormais piloter le projet"),
    ])

    add_heading(doc, "C. Superviser le portefeuille complet", level=2)
    add_step_table(doc, [
        (1, "DSI", "Aller sur le Dashboard"),
        (2, "DSI", "Voir tous les projets (toutes directions)"),
        (3, "DSI", "Filtrer par : statut, phase, direction, chef de projet"),
        (4, "DSI", "Consulter les KPIs : taux de respect des delais, budget, qualite, REP"),
        (5, "DSI", "Exporter en Excel ou PDF"),
    ])

    add_heading(doc, "D. Valider la charte projet", level=2)
    add_step_table(doc, [
        (1, "Systeme", "Notification : charte generee, signature DSI requise"),
        (2, "DSI",     "Ouvrir le projet, onglet Planification"),
        (3, "DSI",     "Consulter la charte PDF"),
        (4, "DSI",     "Valider ou rejeter la charte"),
        (5, "Systeme", "Si validee : projet peut passer en Execution"),
    ])

    add_heading(doc, "E. Valider les livrables et les phases cles", level=2)
    add_step_table(doc, [
        (1, "DSI", "Ouvrir un projet"),
        (2, "DSI", "Aller dans l'onglet Planification ou Execution"),
        (3, "DSI", "Valider ou rejeter les livrables soumis"),
        (4, "DSI", "Acceder a l'historique complet des actions"),
    ])

    add_page_break(doc)

    # ─────────────────────────────────────────────────────────────────────────
    # ROLE : RESPONSABLE SOLUTIONS IT
    # ─────────────────────────────────────────────────────────────────────────
    add_heading(doc, "5. Role : Responsable Solutions IT (Chef de Projet)")
    add_role_header(doc, "RESPONSABLE SOLUTIONS IT",
        "[RSI]",
        "Pilote les projets IT. Acces complet a tous les projets du portefeuille.",
        "Gerer toutes les phases de tous les projets, valider livrables, exporter rapports.")

    add_heading(doc, "A. Gerer le portefeuille de projets", level=2)
    add_step_table(doc, [
        (1, "RSI", "Acceder au Portefeuille ou au Dashboard"),
        (2, "RSI", "Voir tous les projets de l'entreprise"),
        (3, "RSI", "Filtrer et trier selon les besoins"),
        (4, "RSI", "Intervenir sur n'importe quel projet comme Chef de Projet"),
    ])

    add_heading(doc, "B. Piloter un projet (toutes phases)", level=2)
    add_para(doc, "Le Responsable Solutions IT a les memes droits qu'un Chef de Projet sur tous les projets. Voir la section 7 (Chef de Projet) pour le detail complet.")

    add_heading(doc, "C. Valider les livrables", level=2)
    add_step_table(doc, [
        (1, "RSI", "Ouvrir un projet"),
        (2, "RSI", "Aller dans l'onglet correspondant a la phase"),
        (3, "RSI", "Cliquer sur 'Valider' ou 'Rejeter' pour chaque livrable"),
        (4, "RSI", "Ajouter un commentaire si rejet"),
    ])

    add_heading(doc, "D. Generer les rapports", level=2)
    add_step_table(doc, [
        (1, "RSI", "Ouvrir un projet"),
        (2, "RSI", "Exporter la charte projet en PDF"),
        (3, "RSI", "Exporter la fiche projet en Word"),
        (4, "RSI", "Exporter le rapport de cloture en PDF"),
        (5, "RSI", "Exporter le portefeuille en Excel"),
    ])

    add_page_break(doc)

    # ─────────────────────────────────────────────────────────────────────────
    # ROLE : ADMINIT
    # ─────────────────────────────────────────────────────────────────────────
    add_heading(doc, "6. Role : AdminIT")
    add_role_header(doc, "ADMINIT — ADMINISTRATEUR APPLICATIF",
        "[A]",
        "Gere les utilisateurs, les directions et la configuration de l'application.",
        "Administration complete : utilisateurs, roles, directions, services, parametres.")

    add_heading(doc, "A. Creer un utilisateur", level=2)
    add_step_table(doc, [
        (1, "AdminIT", "Aller dans 'Administration > Utilisateurs'"),
        (2, "AdminIT", "Cliquer sur 'Nouvel utilisateur'"),
        (3, "AdminIT", "Remplir : nom, prenom, email, matricule, direction, role"),
        (4, "AdminIT", "Enregistrer"),
        (5, "Systeme",  "Lien d'activation securise genere et envoye par email"),
    ])

    add_heading(doc, "B. Modifier ou desactiver un utilisateur", level=2)
    add_step_table(doc, [
        (1, "AdminIT", "Aller dans 'Administration > Utilisateurs'"),
        (2, "AdminIT", "Rechercher l'utilisateur"),
        (3, "AdminIT", "Modifier : role, direction, service, email"),
        (4, "AdminIT", "Ou desactiver le compte (EstSupprime = true)"),
    ])

    add_heading(doc, "C. Reinitialiser un mot de passe", level=2)
    add_step_table(doc, [
        (1, "AdminIT", "Ouvrir la fiche de l'utilisateur"),
        (2, "AdminIT", "Cliquer sur 'Reinitialiser le mot de passe'"),
        (3, "Systeme",  "Nouveau mot de passe genere et envoye par email a l'utilisateur"),
    ])

    add_heading(doc, "D. Gerer les directions et services", level=2)
    add_step_table(doc, [
        (1, "AdminIT", "Aller dans 'Administration > Directions'"),
        (2, "AdminIT", "Creer / modifier / supprimer une direction"),
        (3, "AdminIT", "Ajouter des services au sein d'une direction"),
        (4, "AdminIT", "Affecter des utilisateurs a une direction"),
    ])

    add_heading(doc, "E. Traiter les demandes de creation de compte", level=2)
    add_step_table(doc, [
        (1, "Systeme",  "Notification : nouvelle demande de compte d'un collaborateur"),
        (2, "AdminIT", "Aller dans 'Administration > Demandes de compte'"),
        (3, "AdminIT", "Consulter la demande (nom, direction, role souhaite)"),
        (4, "AdminIT", "Valider => creer le compte automatiquement"),
        (5, "Systeme",  "Email envoye avec les credentials"),
    ])

    add_page_break(doc)

    # ─────────────────────────────────────────────────────────────────────────
    # ROLE : CHEF DE PROJET
    # ─────────────────────────────────────────────────────────────────────────
    add_heading(doc, "7. Role : Chef de Projet")
    add_role_header(doc, "CHEF DE PROJET",
        "[CP]",
        "Pilote son projet affecte. Responsable de toutes les phases de son projet.",
        "Son projet uniquement : toutes les phases, tous les onglets, generation des livrables.")

    add_heading(doc, "A. Prendre en charge un projet affecte", level=2)
    add_step_table(doc, [
        (1, "Systeme",     "Email de notification : 'Vous etes affecte au projet [Code]'"),
        (2, "Chef Projet", "Se connecter et ouvrir son projet"),
        (3, "Chef Projet", "Consulter le contexte : demande initiale, objectifs, budget alloue"),
        (4, "Systeme",     "Prise en charge enregistree dans le journal d'audit"),
    ])

    add_heading(doc, "B. Phase 1 : Analyse & Clarification", level=2)
    add_step_table(doc, [
        (1, "Chef Projet", "Aller dans l'onglet 'Analyse'"),
        (2, "Chef Projet", "Rediger la note de cadrage"),
        (3, "Chef Projet", "Uploader les documents d'analyse"),
        (4, "Chef Projet", "Ajouter les comptes-rendus de reunion"),
        (5, "DSI/RSI",     "Valider les livrables d'analyse"),
        (6, "Chef Projet", "Demander le passage en Phase 2"),
    ])

    add_heading(doc, "C. Phase 2 : Planification & Validation", level=2)
    add_step_table(doc, [
        (1, "Chef Projet", "Aller dans l'onglet 'Planification'"),
        (2, "Chef Projet", "Creer le planning : ajouter des taches (titre, responsable, date debut/fin)"),
        (3, "Chef Projet", "Definir le budget previsionnel"),
        (4, "Chef Projet", "Affecter les ressources humaines (profil, TJM)"),
        (5, "Chef Projet", "Generer la charte projet en PDF"),
        (6, "Chef Projet", "Initialiser le dossier de signature electronique"),
        (7, "Chef Projet", "Signer la charte (sa signature + DM + DSI)"),
        (8, "Systeme",     "CharteValidee = true quand toutes les signatures sont presentes"),
        (9, "Chef Projet", "Demander le passage en Phase 3"),
    ])

    add_heading(doc, "D. Phase 3 : Execution & Suivi", level=2)
    add_step_table(doc, [
        (1, "Chef Projet", "Aller dans l'onglet 'Execution'"),
        (2, "Chef Projet", "Mettre a jour l'avancement des taches (%)"),
        (3, "Chef Projet", "Ajouter et gerer les risques (probabilite, impact, plan de mitigation)"),
        (4, "Chef Projet", "Uploader les livrables de la phase"),
        (5, "Chef Projet", "Rediger le bilan hebdomadaire (avancement, budget, risques, actions)"),
        (6, "Chef Projet", "Suivre le budget reel vs previsionnel"),
        (7, "Systeme",     "RAG calcule automatiquement : Vert / Ambre / Rouge"),
        (8, "Chef Projet", "Configurer la collaboration Teams/Planner si besoin"),
        (9, "Chef Projet", "Demander le passage en Phase 4 (UAT)"),
    ])

    add_heading(doc, "E. Phase 4 : UAT & MEP", level=2)
    add_step_table(doc, [
        (1, "Chef Projet", "Aller dans l'onglet 'UAT'"),
        (2, "Chef Projet", "Creer une campagne de test (nom, environnement, date)"),
        (3, "Chef Projet", "Ajouter les cas de test (titre, description, resultat attendu, priorite)"),
        (4, "Chef Projet", "Executer les cas de test (Reussi / Echec / Bloque / Non applicable)"),
        (5, "Chef Projet", "Gerer les anomalies detectees"),
        (6, "DM",          "Valider la recette (RecetteValidee = true)"),
        (7, "Chef Projet", "Marquer la MEP effectuee avec la date reelle"),
        (8, "Chef Projet", "Cliquer sur 'Fin UAT' pour passer en Phase 5"),
    ])

    add_heading(doc, "F. Phase 5 : Cloture & Lecons apprises", level=2)
    add_step_table(doc, [
        (1, "Chef Projet", "Aller dans l'onglet 'Cloture'"),
        (2, "Chef Projet", "Rediger le rapport de cloture complet"),
        (3, "Chef Projet", "Saisir les lecons apprises"),
        (4, "Chef Projet", "Donner une note de satisfaction globale"),
        (5, "Chef Projet", "Exporter le rapport de cloture en PDF"),
        (6, "Chef Projet", "Soumettre la demande de cloture"),
        (7, "DM/Demandeur","Valider la cloture"),
        (8, "Systeme",     "Projet passe en statut 'Cloture'"),
    ])

    add_page_break(doc)

    # ─────────────────────────────────────────────────────────────────────────
    # MATRICE DES DROITS
    # ─────────────────────────────────────────────────────────────────────────
    add_heading(doc, "8. Matrice des droits par role")

    fonctionnalites = [
        ("Soumettre une demande de projet",    "X", "",  "",  "",  "",  ""),
        ("Voir ses propres demandes",          "X", "X", "X", "X", "X", "X"),
        ("Valider/rejeter demandes (DM)",      "",  "X", "X", "",  "",  "X"),
        ("Valider/rejeter demandes (DSI)",     "",  "",  "X", "X", "",  "X"),
        ("Creer un projet",                    "",  "",  "X", "X", "",  "X"),
        ("Affecter un Chef de Projet",         "",  "",  "X", "X", "",  "X"),
        ("Voir tous les projets",              "",  "",  "X", "X", "",  "X"),
        ("Voir projets de sa direction",       "",  "X", "",  "",  "",  ""),
        ("Voir son projet (lecture seule)",    "X", "",  "",  "",  "",  ""),
        ("Piloter un projet (toutes phases)",  "",  "",  "",  "X", "X", "X"),
        ("Gerer les taches",                   "",  "",  "",  "X", "X", "X"),
        ("Gerer les risques",                  "",  "",  "",  "X", "X", "X"),
        ("Uploader les livrables",             "",  "",  "",  "X", "X", "X"),
        ("Valider les livrables",              "",  "",  "X", "X", "",  "X"),
        ("Generer la charte PDF",              "",  "",  "",  "X", "X", "X"),
        ("Signer la charte",                   "",  "X", "X", "",  "X", "X"),
        ("Creer des cas de test",              "",  "",  "",  "X", "X", "X"),
        ("Executer des cas de test",           "",  "X", "X", "X", "X", "X"),
        ("Valider la recette (DM)",            "",  "X", "",  "",  "",  ""),
        ("Configurer collaboration Teams",     "",  "",  "",  "X", "X", "X"),
        ("Rediger bilan hebdomadaire",         "",  "",  "",  "X", "X", "X"),
        ("Rapport de cloture",                 "",  "",  "",  "X", "X", "X"),
        ("Valider la cloture",                 "X", "X", "X", "",  "",  ""),
        ("Dashboard & KPIs",                   "",  "",  "X", "X", "",  "X"),
        ("Exporter Excel / PDF",               "",  "",  "X", "X", "X", "X"),
        ("Gerer les utilisateurs",             "",  "",  "",  "",  "",  "X"),
        ("Gerer directions & services",        "",  "",  "",  "",  "",  "X"),
        ("Parametres systeme",                 "",  "",  "",  "",  "",  "X"),
        ("Journal d'audit",                    "",  "",  "X", "X", "X", "X"),
    ]

    roles_cols = ["Fonctionnalite", "Demandeur", "DM/Sponsor", "DSI", "Resp. Sol. IT", "Chef Projet", "AdminIT"]
    t = doc.add_table(rows=len(fonctionnalites)+1, cols=len(roles_cols))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER

    col_w = [Cm(6.5), Cm(2), Cm(2.2), Cm(1.8), Cm(2.5), Cm(2.2), Cm(1.8)]
    for ci, (h, w) in enumerate(zip(roles_cols, col_w)):
        c = t.rows[0].cells[ci]
        c.width = w
        set_cell_bg(c, NAVY_HEX)
        set_cell_borders(c, NAVY_HEX)
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.font.bold = True
        r.font.color.rgb = WHITE
        r.font.size = Pt(9)
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for i, row_data in enumerate(fonctionnalites):
        row = t.rows[i+1]
        row.height = Cm(0.6)
        bg_base = LIGHT_HEX if i % 2 == 0 else WHITE_HEX
        for ci, val in enumerate(row_data):
            c = row.cells[ci]
            if ci == 0:
                set_cell_bg(c, bg_base)
                set_cell_borders(c, "CCCCCC")
                p = c.paragraphs[0]
                r = p.add_run(val)
                r.font.size = Pt(9)
                r.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
            else:
                if val == "X":
                    set_cell_bg(c, "D4EDDA")
                    set_cell_borders(c, "28A745")
                    p = c.paragraphs[0]
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    r = p.add_run("OK")
                    r.font.bold = True
                    r.font.size = Pt(8)
                    r.font.color.rgb = RGBColor(0x15, 0x52, 0x24)
                else:
                    set_cell_bg(c, bg_base)
                    set_cell_borders(c, "CCCCCC")
                    p = c.paragraphs[0]
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    r = p.add_run("—")
                    r.font.size = Pt(8)
                    r.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    add_page_break(doc)

    # ─────────────────────────────────────────────────────────────────────────
    # WORKFLOW GLOBAL SYNTHETIQUE
    # ─────────────────────────────────────────────────────────────────────────
    add_heading(doc, "9. Workflow global — Vue synthetique")

    workflow_global = [
        ("Demandeur",      "Soumet une demande de projet avec tous les details"),
        ("Systeme",        "Email automatique envoye au Directeur Metier"),
        ("DM",             "Valide / Rejette / Renvoie la demande"),
        ("Systeme",        "Si validee DM : email envoye a la DSI"),
        ("DSI",            "Valide / Rejette la demande"),
        ("Systeme",        "Si validee DSI : creation automatique du projet (code unique)"),
        ("DSI",            "Affecte un Responsable Solutions IT (Chef de Projet)"),
        ("Chef de Projet", "Phase 1 — Analyse : note de cadrage, livrables, validation"),
        ("Chef de Projet", "Phase 2 — Planification : planning, budget, ressources, charte signee"),
        ("Chef de Projet", "Phase 3 — Execution : taches, risques, livrables, bilans hebdo, RAG"),
        ("Chef de Projet", "Phase 4 — UAT : campagnes de test, cas de test, execution, anomalies"),
        ("DM",             "Valide la recette (GO MEP)"),
        ("Chef de Projet", "Marque la MEP effectuee — passe le projet en Phase 5"),
        ("Chef de Projet", "Phase 5 — Cloture : rapport, lecons apprises, demande de cloture"),
        ("DM / Demandeur", "Valident la cloture"),
        ("Systeme",        "Projet passe en statut 'Cloture' — archive dans le portefeuille"),
    ]

    add_step_table(doc, [(i+1, a, d) for i, (a, d) in enumerate(workflow_global)])

    add_info_box(doc,
        "Notifications automatiques",
        "A chaque etape cle, l'application envoie automatiquement un email aux acteurs concernes.\n"
        "Les notifications Teams sont aussi disponibles si l'integration Microsoft 365 est configuree.",
        LIGHT_HEX, GOLD_HEX)

    add_info_box(doc,
        "Indicateurs RAG — calcul automatique",
        "Vert  : Avancement normal, budget respecte, pas de risque majeur\n"
        "Ambre : Retard modere (<15%), depassement budget (<10%), risques matrises\n"
        "Rouge : Retard important (>15%), depassement budget (>10%), risque critique",
        LIGHT_HEX, BLUE_HEX)

    return doc

if __name__ == "__main__":
    doc = build_document()
    out = "Workflow_Par_Role_GestionProjetsIT.docx"
    doc.save(out)
    print(f"OK Fichier Word genere : {out}")
