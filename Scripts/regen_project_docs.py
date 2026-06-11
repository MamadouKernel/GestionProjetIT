from __future__ import annotations

from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Iterable, Sequence

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(r"C:\Users\kerne\Downloads\aldric\projets\GestionProjetIT")
LOGO = ROOT / "wwwroot" / "images" / "LOGO_COTE_D_IVOIRE_TERMINAL.png"
WORKFLOW_DOC = ROOT / "Workflow_Par_Role_GestionProjetsIT.docx"
FULL_DOC = ROOT / "Documentation_Complete_Application_GestionProjetsIT.docx"

PRIMARY = RGBColor(28, 92, 156)
SECONDARY = RGBColor(89, 102, 122)
ACCENT = RGBColor(230, 244, 255)
LIGHT = RGBColor(245, 247, 250)


def set_cell_shading(cell, hex_fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_fill)
    tc_pr.append(shd)


def set_table_borders(table) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "D9E1EA")
        borders.append(element)
    tbl_pr.append(borders)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def configure_document(doc: Document, header_label: str) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.08

    for name, size, color in (
        ("Title", 22, PRIMARY),
        ("Heading 1", 16, PRIMARY),
        ("Heading 2", 13, PRIMARY),
        ("Heading 3", 11, SECONDARY),
    ):
        style = styles[name]
        style.font.name = "Arial"
        style.font.bold = True
        style.font.size = Pt(size)
        style.font.color.rgb = color

    header = section.header
    header_p = header.paragraphs[0]
    header_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if LOGO.exists():
        header_p.add_run().add_picture(str(LOGO), width=Inches(0.45))
        header_p.add_run("  ")
    run = header_p.add_run(header_label)
    run.font.name = "Arial"
    run.font.size = Pt(8.5)
    run.font.color.rgb = SECONDARY

    footer = section.footer
    footer_p = footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_run = footer_p.add_run(f"Côte d'Ivoire Terminal • Gestion Projets IT • {datetime.now():%d/%m/%Y}")
    footer_run.font.name = "Arial"
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = SECONDARY


def add_cover(doc: Document, title: str, subtitle: str, purpose: str) -> None:
    if LOGO.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(LOGO), width=Inches(1.1))

    title_p = doc.add_paragraph(style="Title")
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.add_run(title)

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub_p.add_run(subtitle)
    r.font.name = "Arial"
    r.font.size = Pt(12)
    r.font.color.rgb = SECONDARY

    doc.add_paragraph("")

    purpose_p = doc.add_paragraph()
    purpose_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = purpose_p.add_run(purpose)
    run.font.name = "Arial"
    run.font.size = Pt(11)

    meta = doc.add_table(rows=3, cols=2)
    meta.style = "Table Grid"
    set_table_borders(meta)
    meta.cell(0, 0).text = "Version"
    meta.cell(0, 1).text = "Mise à jour 2026-05-21"
    meta.cell(1, 0).text = "Périmètre"
    meta.cell(1, 1).text = "Application Gestion Projets IT"
    meta.cell(2, 0).text = "Objet"
    meta.cell(2, 1).text = "Référentiel fonctionnel, rôles, écrans et mode de renseignement"
    for row in meta.rows:
        row.cells[0].width = Cm(4.2)
        row.cells[1].width = Cm(11.8)
        row.cells[0].paragraphs[0].runs[0].font.bold = True
        set_cell_shading(row.cells[0], "EAF3FF")
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    doc.add_page_break()


def add_section_intro(doc: Document, title: str, text: str) -> None:
    doc.add_heading(title, level=1)
    doc.add_paragraph(text)


def add_bullets(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_table(doc: Document, headers: Sequence[str], rows: Sequence[Sequence[str]], widths_cm: Sequence[float] | None = None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_borders(table)

    header = table.rows[0]
    set_repeat_table_header(header)
    for idx, text in enumerate(headers):
        cell = header.cells[idx]
        cell.text = text
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_shading(cell, "DDEBF7")
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.runs[0]
        run.font.bold = True
        run.font.name = "Arial"
        run.font.size = Pt(9.5)

    for row_values in rows:
        row = table.add_row()
        for idx, value in enumerate(row_values):
            cell = row.cells[idx]
            cell.text = value
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(9.5)

    if widths_cm:
        for row in table.rows:
            for idx, width in enumerate(widths_cm):
                row.cells[idx].width = Cm(width)


@dataclass
class ScreenRow:
    screen: str
    actors: str
    behavior: str
    how_to_fill: str


def add_screen_matrix(doc: Document, title: str, intro: str, rows: Sequence[ScreenRow]) -> None:
    doc.add_heading(title, level=2)
    doc.add_paragraph(intro)
    add_table(
        doc,
        ["Écran", "Acteurs", "Comportement attendu", "Comment renseigner / utiliser"],
        [[r.screen, r.actors, r.behavior, r.how_to_fill] for r in rows],
        widths_cm=[4.1, 3.0, 5.8, 5.8],
    )


def workflow_doc() -> Document:
    doc = Document()
    configure_document(doc, "Workflow par rôle • Gestion Projets IT")
    add_cover(
        doc,
        "Workflow par rôle",
        "Gestion Projets IT - Côte d'Ivoire Terminal",
        "Document de référence pour comprendre qui agit, sur quel écran, à quel moment, et quelles données doivent être renseignées pour faire avancer le projet.",
    )

    add_section_intro(
        doc,
        "1. Logique générale de l'application",
        "L'application pilote un cycle complet : demande, analyse, planification, exécution, UAT/MEP, clôture. "
        "Chaque étape est portée par des écrans dédiés, des livrables, des validations et des règles de blocage automatiques. "
        "La règle d'usage est simple : la saisie structurée se fait directement dans l'application, les pièces externes ou signées sont déposées comme livrables.",
    )
    add_bullets(doc, [
        "Le Demandeur exprime le besoin et suit son dossier.",
        "Le Directeur Métier valide le besoin puis la charte et la planification côté métier.",
        "Le Chef de Projet prépare les contenus opérationnels et tient le projet à jour.",
        "La DSI ou le RSIT délégué valide les jalons IT et autorise le passage entre phases.",
        "L'Admin IT paramètre l'application et conserve une capacité d'accès transverse.",
    ])

    add_screen_matrix(doc, "2. Écrans transverses", "Écrans disponibles pour plusieurs rôles.", [
        ScreenRow("Tableau de bord", "Tous selon le rôle", "Affiche les indicateurs, validations en attente et alertes utiles au rôle connecté.", "Aucune saisie ; utiliser pour repérer les points d'action prioritaires."),
        ScreenRow("Notifications", "Tous", "Centralise les actions attendues, retours et mises à jour système.", "Lire, marquer comme lu et ouvrir les liens vers l'écran concerné."),
        ScreenRow("Historique & traçabilité", "CP, DM, DSI, RSIT, AdminIT", "Expose la chronologie des décisions, validations et événements projet.", "Aucune saisie ; utiliser comme piste d'audit et de contrôle."),
        ScreenRow("Mon profil", "Tous", "Affiche identité, rôle, sécurité et données personnelles de l'utilisateur.", "Mettre à jour uniquement les champs autorisés, notamment le mot de passe selon la politique de sécurité."),
    ])

    add_section_intro(doc, "3. Workflow par phase", "La bonne compréhension du cycle projet passe par les écrans et prérequis ci-dessous.")

    add_screen_matrix(doc, "3.1 Demande initiale", "La phase de demande prépare le besoin avant création du projet.", [
        ScreenRow("Mes demandes", "Demandeur", "Liste les demandes créées avec leurs statuts et actions disponibles.", "Consulter l'état du dossier, corriger une demande si elle revient en révision, ouvrir le détail pour suivre les validations."),
        ScreenRow("Nouvelle demande", "Demandeur", "Permet de saisir le besoin, les objectifs, le contexte et le premier dossier documentaire.", "Renseigner le besoin métier, l'objectif, le contexte, les bénéfices attendus et joindre le cahier des charges initial."),
        ScreenRow("Validations DM", "Directeur Métier", "Permet de valider, rejeter ou renvoyer les demandes de son périmètre.", "Contrôler la cohérence métier, commenter en cas de refus et valider uniquement les demandes suffisamment cadrées."),
        ScreenRow("Validations DSI", "DSI / RSIT délégué", "Décision IT de création du projet après validation DM.", "Vérifier faisabilité, alignement IT et disponibilité de prise en charge avant validation."),
    ])

    add_screen_matrix(doc, "3.2 Analyse & Clarification", "La phase Analyse prépare la charte et les livrables de cadrage.", [
        ScreenRow("Détail projet - onglet Analyse", "Chef de Projet", "Écran de travail principal : contexte, objectifs, documents d'analyse, risques, charte et passage à la planification.", "Mettre à jour les informations issues du besoin, charger le cahier d'analyse technique et la note de cadrage, gérer les risques, générer la charte PDF."),
        ScreenRow("Charte Projet", "Chef de Projet, DSI, RSIT, AdminIT, DM en consultation", "Écran dédié à l'édition de la charte, au dépôt de la version signée et à l'enregistrement des signatures.", "Renseigner la charte, déposer la version signée, cocher séparément la signature Sponsor/DM et la signature Chef de Projet, enregistrer les signatures même sans re-upload."),
        ScreenRow("Validations Projet", "DM puis DSI / RSIT délégué", "Validation séquentielle de la charte avant passage en planification.", "Le DM valide après dépôt de la charte signée complète ; la DSI/RSIT valide ensuite. Si le dossier est incomplet, la vue l'indique explicitement."),
        ScreenRow("Passage Analyse -> Planification", "Chef de Projet", "Le bouton de validation de phase ne s'active que si la charte PDF existe, la charte signée est complète et les validations DM/DSI sont faites.", "Cliquer sur 'Valider la phase Analyse' seulement quand l'alerte ne signale plus aucun élément manquant."),
    ])

    add_screen_matrix(doc, "3.3 Planification & Validation", "La planification est désormais majoritairement native.", [
        ScreenRow("Détail projet - onglet Planification", "Chef de Projet", "Écran central pour construire le planning, la RACI, la communication, le budget et le PV de kick-off.", "Saisir d'abord les tâches dans le Gantt, puis enregistrer la planification. Compléter ensuite la RACI, le plan de communication, le budget et le PV de kick-off natifs."),
        ScreenRow("Planning interactif / Gantt", "Chef de Projet", "Permet d'ajouter les tâches, jalons, dates, responsables, dépendances et commentaires.", "Cliquer 'Ajouter une tâche', renseigner WBS, tâche, début, fin et responsable au minimum, puis enregistrer. Le bouton ne sauvegarde pas à lui seul."),
        ScreenRow("Générer Planning + WBS", "Chef de Projet", "Transforme les données natives en livrables officiels brandés avec logo.", "Utiliser uniquement après avoir enregistré au moins une tâche. Le système produit Planning détaillé et WBS au format Excel."),
        ScreenRow("Validation Directeur Métier", "Directeur Métier", "Feu vert métier de la planification.", "Vérifier cohérence planning, budget, RACI et kick-off avant validation."),
        ScreenRow("Validation DSI", "DSI / RSIT délégué", "Autorise le passage en exécution après validation DM et présence des livrables obligatoires.", "Contrôler la complétude des livrables, la charge et la gouvernance IT puis cliquer sur le passage en exécution."),
    ])

    add_screen_matrix(doc, "3.4 Exécution", "La phase Exécution suit la réalisation projet.", [
        ScreenRow("Détail projet - onglet Exécution", "Chef de Projet", "Sert à suivre l'avancement, les actions réalisées, les blocages, les décisions et les écarts.", "Renseigner les faits marquants, commentaires d'avancement, jalons réels, blocages, décisions et synthèse de charge."),
        ScreenRow("Charges / Capacité", "Chef de Projet, DSI, AdminIT", "Module de saisie et de pilotage des charges prévues/réelles.", "Saisir les heures par ressource et par semaine, commenter les écarts, surveiller les surcharges et la disponibilité."),
        ScreenRow("Risques", "Chef de Projet", "Permet d'ajouter et suivre les risques projet.", "Créer un risque avec description, probabilité, impact, responsable et plan de mitigation ; maintenir ces données à jour pendant l'exécution."),
    ])

    add_screen_matrix(doc, "3.5 UAT & MEP", "La phase UAT/MEP prépare la recette et la mise en production.", [
        ScreenRow("Détail projet - onglet UAT & MEP", "Chef de Projet, DSI", "Centralise recette, anomalies, plan de mise en production et hypercare.", "Renseigner dates de recette, périmètre testé, prérequis MEP, plan de rollback, incidents et statut hypercare."),
        ScreenRow("Cas de test", "Chef de Projet / équipe projet", "Suit les cas de test et l'état de la recette.", "Créer ou mettre à jour les cas de test, indiquer résultat, commentaire et anomalie associée."),
    ])

    add_screen_matrix(doc, "3.6 Clôture", "La clôture est une validation multi-acteurs.", [
        ScreenRow("Détail projet - onglet Clôture", "Chef de Projet", "Prépare le bilan final, le transfert au run et les leçons apprises.", "Renseigner statut final, bilan, documentation run, support informé, exploitation prête et commentaires de clôture."),
        ScreenRow("Mes validations clôture", "Demandeur", "Le demandeur confirme que le résultat livré est conforme au besoin.", "Relire la clôture, valider ou rejeter avec commentaire."),
        ScreenRow("Validations Clôture DM", "Directeur Métier", "Validation métier finale du projet.", "Contrôler la valeur métier délivrée et les impacts résiduels avant validation."),
        ScreenRow("Validations Clôture DSI", "DSI / AdminIT", "Dernier feu vert avant passage en projet clôturé.", "Valider le bilan IT, le transfert run et la complétude documentaire."),
    ])

    add_section_intro(doc, "4. Workflow par rôle", "Les tableaux suivants donnent la logique d'usage cible par profil.")
    add_table(doc,
              ["Rôle", "Écrans principaux", "Ce que le rôle fait", "Règles de comportement"],
              [
                  ["Demandeur", "Mes demandes, Nouvelle demande, Détails, validations clôture", "Exprime le besoin, suit le dossier, répond aux retours et valide la clôture.", "Ne modifie pas la planification ni les validations IT ; agit sur son besoin et son retour utilisateur."],
                  ["Directeur Métier", "Validations DM, Validations Projet, Validations Clôture DM, Projets de ma direction", "Valide la demande, la charte, la planification et la clôture métier.", "Ne construit pas les données projet ; vérifie, arbitre et commente."],
                  ["Chef de Projet", "Détail projet, Charte Projet, Planification, Exécution, UAT & MEP, Clôture, Charges", "Prépare et maintient tout le dossier opérationnel du projet.", "Renseigne en continu les écrans de travail ; génère les livrables natifs et alerte les validateurs."],
                  ["DSI", "Validations DSI, Validations Projet, Planification, Dashboard, clôture DSI", "Valide les jalons IT et autorise les passages de phase.", "N'intervient qu'après les prérequis métier/documentaires ; bloque si la gouvernance n'est pas satisfaisante."],
                  ["RSIT", "Portefeuille, Validations DSI si délégation, consultation projet", "Supporte la DSI et agit comme délégué sur les validations autorisées.", "Ne valide que si la délégation DSI est active."],
                  ["AdminIT", "Administration complète, Autorisations / Droits, tous les modules", "Paramètre l'application, gère les droits et peut intervenir en secours.", "Doit préserver la cohérence des workflows et éviter de contourner les validations métier sans justification."],
              ],
              widths_cm=[2.6, 5.0, 5.1, 5.3])

    add_section_intro(doc, "5. Règles de blocage importantes", "Les écrans appliquent des contrôles automatiques pour empêcher les passages de phase prématurés.")
    add_bullets(doc, [
        "Analyse -> Planification : nécessite la charte PDF, la charte signée complète, la validation DM et la validation DSI.",
        "Planification -> Exécution : nécessite Planning détaillé, WBS, Matrice RACI, Schéma de communication, Budget prévisionnel, PV de kick-off et les validations DM puis DSI.",
        "DSI / RSIT ne peut pas valider la charte si le Directeur Métier n'a pas déjà validé.",
        "Les documents générés nativement par l'application portent le logo CIT et remplacent les versions générées précédentes du même type.",
    ])

    return doc


def full_doc() -> Document:
    doc = Document()
    configure_document(doc, "Documentation complète • Gestion Projets IT")
    add_cover(
        doc,
        "Documentation complète de l'application",
        "Gestion Projets IT - Côte d'Ivoire Terminal",
        "Manuel fonctionnel et d'exploitation décrivant le comportement attendu de chaque écran principal et la manière correcte de le renseigner.",
    )

    add_section_intro(doc, "1. Objet du document", "Cette documentation décrit le fonctionnement de l'application, module par module, en insistant sur les écrans réellement utilisés, les comportements attendus, les données à saisir et les règles de blocage.")

    add_section_intro(doc, "2. Navigation générale", "L'utilisateur travaille depuis une coque commune : menu latéral, top bar, notifications, profil et dashboard. Les écrans visibles dépendent du rôle et de la matrice d'autorisations.")
    add_bullets(doc, [
        "Le menu latéral n'affiche que les entrées autorisées pour le rôle connecté.",
        "Les badges, alertes et notifications servent de système de priorisation.",
        "Les boutons d'action en icône seule ouvrent généralement le détail, la validation ou la suppression selon leur couleur.",
        "Un écran de détail projet est piloté par onglets ; chaque onglet correspond à une phase ou à un sous-module métier.",
    ])

    add_screen_matrix(doc, "3. Module Demandes", "Écrans du cycle de demande avant création du projet.", [
        ScreenRow("Mes demandes", "Demandeur", "Affiche la liste des demandes créées et leur statut.", "Utiliser comme point d'entrée principal pour suivre le traitement et rouvrir un dossier."),
        ScreenRow("Nouvelle demande", "Demandeur", "Permet de saisir le besoin initial et les premiers justificatifs.", "Renseigner le contexte, les objectifs, les bénéfices attendus, le périmètre et joindre le cahier des charges ou annexes."),
        ScreenRow("Détail demande", "Demandeur, DM, DSI", "Présente la demande, ses pièces et ses décisions.", "Ajouter un commentaire ou des documents complémentaires si le workflow le demande."),
        ScreenRow("Validations DM", "Directeur Métier", "Valide ou rejette la demande côté métier.", "Contrôler le besoin, le sponsor, la direction, la valeur métier et commenter toute décision."),
        ScreenRow("Validations DSI", "DSI / RSIT délégué", "Décision de faisabilité IT et de création de projet.", "Valider uniquement si le dossier est cadré et réaliste en charge / ressources."),
        ScreenRow("Historique Actions DM", "DM, DSI, AdminIT", "Historique des actions et commentaires métier.", "Aucune saisie ; exploiter pour relecture et audit."),
    ])

    add_screen_matrix(doc, "4. Détail projet - onglet Synthèse", "Vue de référence du projet.", [
        ScreenRow("Synthèse", "CP, DM, DSI, RSIT, AdminIT", "Présente l'identité projet, la progression, les validations, les risques et la vue d'ensemble.", "Relire régulièrement pour vérifier les voyants, le RAG, l'avancement et les dernières alertes."),
    ])

    add_screen_matrix(doc, "5. Détail projet - onglet Analyse", "Phase de cadrage détaillé et de préparation de la charte.", [
        ScreenRow("Contexte / Objectifs", "Chef de Projet", "Affiche les informations héritées de la demande et permet leur maintien.", "Vérifier la cohérence avec la demande initiale, reformuler si nécessaire de façon opérationnelle."),
        ScreenRow("Documents d'analyse obligatoires", "Chef de Projet", "Zone des livrables 'Cahier d'analyse technique' et 'Note de cadrage'.", "Déposer le bon document dans le bon slot. Le type de livrable choisi à l'upload détermine l'affichage."),
        ScreenRow("Autres livrables d'analyse", "Chef de Projet", "Zone pour études, annexes, comptes rendus ou supports complémentaires.", "Ajouter les pièces utiles au cadrage mais non obligatoires."),
        ScreenRow("Générer PDF Charte", "Chef de Projet", "Crée le livrable officiel 'Charte projet'.", "Cliquer après avoir consolidé les données de charte. Le PDF généré devient un livrable du projet."),
        ScreenRow("Valider la phase Analyse", "Chef de Projet", "Passe à la planification si toutes les conditions sont remplies.", "Ne cliquer qu'après disparition des alertes rouges et jaunes de complétude."),
    ])

    add_screen_matrix(doc, "6. Écran Charte Projet", "Sous-module de la charte.", [
        ScreenRow("Édition de la charte", "Chef de Projet / DSI / AdminIT", "Permet de compléter les champs structurés de la charte.", "Renseigner objectif, périmètre, contraintes, risques initiaux, acteurs et jalons."),
        ScreenRow("Télécharger Word / PDF", "CP, DSI, AdminIT", "Génère les versions bureautiques de la charte.", "Utiliser pour partager la version officielle à faire signer."),
        ScreenRow("Version signée", "CP, DSI, RSIT, AdminIT", "Dépôt de la charte signée et enregistrement séparé des signatures.", "Déposer la version signée, cocher Sponsor/DM et Chef de Projet, puis enregistrer les signatures même sans re-upload."),
        ScreenRow("Dossier de signature", "CP, DSI, RSIT, AdminIT", "Prépare ou suit le dossier de signature électronique quand il est utilisé.", "Initialiser, envoyer et suivre les signataires si le processus de signature externe est utilisé."),
    ])

    add_screen_matrix(doc, "7. Écran Validations Projet", "Validation de la charte par les décideurs.", [
        ScreenRow("Carte projet à valider", "DM / DSI / RSIT délégué", "Affiche le statut de la charte, les signatures manquantes et les actions possibles.", "Le bouton vert ne s'active que si le dossier de charte est complet ; le bouton rouge sert au refus avec commentaire ; l'œil ouvre le détail."),
        ScreenRow("Validation DM", "Directeur Métier", "Premier niveau de validation de la charte.", "Relire la charte et vérifier que le sponsor/DM et le CP ont bien signé avant validation."),
        ScreenRow("Validation DSI", "DSI / RSIT délégué", "Deuxième niveau, après validation DM.", "Valider seulement après confirmation du dossier complet et du feu vert métier."),
    ])

    add_screen_matrix(doc, "8. Détail projet - onglet Planification", "Module natif de planification.", [
        ScreenRow("Checklist des livrables", "CP, DM, DSI", "Montre ce qui est généré ou manquant avant exécution.", "Se servir de cette liste comme check final avant validation DM puis DSI."),
        ScreenRow("Planning interactif / Gantt", "Chef de Projet", "Éditeur natif des tâches et jalons.", "Ajouter une tâche, renseigner WBS, tâche, dates, responsable, dépendances et commentaire, puis enregistrer."),
        ScreenRow("RACI native", "Chef de Projet", "Éditeur des responsabilités projet.", "Créer une ligne par activité ; indiquer R, A, C, I clairement pour chaque acteur."),
        ScreenRow("Plan de communication natif", "Chef de Projet", "Éditeur des instances, fréquences et canaux.", "Définir chaque réunion/instance, son objectif, sa fréquence, ses participants et le responsable."),
        ScreenRow("Budget natif", "Chef de Projet", "Éditeur des lignes budgétaires détaillées.", "Saisir un poste par ligne, description, montant et commentaire ; vérifier le total calculé."),
        ScreenRow("PV de kick-off natif", "Chef de Projet", "Fiche structurée du lancement officiel.", "Renseigner date, heure, lieu, animateur, objectifs, participants, ordre du jour, décisions et actions."),
        ScreenRow("Enregistrer la planification", "Chef de Projet", "Sauvegarde toutes les données natives de planification.", "Toujours enregistrer avant de générer les livrables."),
        ScreenRow("Générer Planning + WBS", "Chef de Projet", "Génère les livrables natifs brandés avec logo.", "Ne fonctionne qu'après enregistrement d'au moins une tâche."),
        ScreenRow("Validation DM / Validation DSI", "DM puis DSI/RSIT", "Valident la planification et le passage en exécution.", "Le DM valide d'abord ; la DSI ne peut valider qu'après et si tous les livrables obligatoires sont présents."),
    ])

    add_screen_matrix(doc, "9. Détail projet - onglet Exécution", "Suivi opérationnel du projet.", [
        ScreenRow("Avancement et commentaires", "Chef de Projet", "Expose l'état réel d'exécution.", "Renseigner dates réelles, écarts, décisions, actions réalisées, actions à venir et blocages."),
        ScreenRow("Suivi des charges", "Chef de Projet / DSI", "Pilote charge prévue et charge réelle.", "Saisir chaque semaine les heures, commentaires et charges par ressource."),
        ScreenRow("Risques et anomalies", "Chef de Projet", "Suit les incidents et risques projet.", "Créer les risques avec plan de mitigation et mettre à jour leur statut."),
    ])

    add_screen_matrix(doc, "10. Détail projet - onglet UAT & MEP", "Préparation de la recette et de la mise en production.", [
        ScreenRow("Recette / anomalies", "Chef de Projet / équipe", "Suit les tests, anomalies et validations recette.", "Décrire le périmètre testé, le résultat, les anomalies et leurs corrections."),
        ScreenRow("Plan de MEP", "Chef de Projet / DSI", "Prépare la bascule et l'hypercare.", "Renseigner prérequis, plan de rollback, incidents, statut hypercare et résultat de MEP."),
    ])

    add_screen_matrix(doc, "11. Détail projet - onglet Collaboration et Historique", "Communication et audit.", [
        ScreenRow("Collaboration", "Équipe projet", "Expose échanges, documents et coordination autour du projet.", "Ajouter les informations utiles à la coopération opérationnelle."),
        ScreenRow("Historique", "Tous rôles autorisés", "Liste chronologique des événements importants.", "Aucune saisie ; exploiter pour comprendre l'évolution du projet."),
    ])

    add_screen_matrix(doc, "12. Détail projet - onglet Clôture", "Fermeture contrôlée du projet.", [
        ScreenRow("Bilan de clôture", "Chef de Projet", "Prépare le statut final et le retour d'expérience.", "Renseigner bilan, écarts, transfert au run, documentation, support et conclusions."),
        ScreenRow("Validations clôture", "Demandeur, DM, DSI", "Validation séquentielle de la fin de projet.", "Chaque acteur valide depuis son écran dédié lorsque son étape est atteinte."),
    ])

    add_screen_matrix(doc, "13. Portefeuille, projets de direction et historiques", "Vues de pilotage transversal.", [
        ScreenRow("Portefeuille DSI", "DSI, RSIT, AdminIT", "Vision consolidée des projets et exports de reporting.", "Filtrer, ouvrir les détails projet, lancer les exports PDF/Excel."),
        ScreenRow("Projets de ma direction", "Directeur Métier", "Vue filtrée des projets de la direction.", "Utiliser pour supervision métier et accès rapide aux validations."),
        ScreenRow("Historique de mes projets / Historique DM", "DM, CP, DSI", "Fiches repliables retraçant le cycle de vie de chaque projet.", "Ouvrir les cartes, lire la timeline et analyser la progression ou les blocages."),
    ])

    add_screen_matrix(doc, "14. Administration", "Écrans de paramétrage et de gouvernance.", [
        ScreenRow("Utilisateurs", "AdminIT", "Gestion des comptes et informations utilisateurs.", "Créer, mettre à jour, réinitialiser et suivre les comptes en respectant la politique de mot de passe."),
        ScreenRow("Rôles", "AdminIT", "Affectation des rôles et gestion de la matrice de droits.", "Attribuer les rôles selon le besoin métier et réinitialiser la matrice quand le standard doit être réappliqué."),
        ScreenRow("Directions / Services", "AdminIT", "Référentiels organisationnels.", "Maintenir les libellés et rattachements organisationnels propres."),
        ScreenRow("Délégations", "AdminIT", "Gestion des délégations DSI et CP si prévues.", "Créer ou arrêter une délégation avec période de validité et acteur remplaçant."),
        ScreenRow("Paramètres", "AdminIT", "Réglages applicatifs et règles globales.", "Modifier uniquement les paramètres maîtrisés et validés côté exploitation."),
        ScreenRow("Autorisations / Droits", "AdminIT", "Matrice de permissions par rôle.", "Activer ou désactiver chaque écran/action par rôle, puis se reconnecter pour constater l'effet."),
    ])

    add_section_intro(doc, "15. Documents et livrables", "L'application gère deux familles de documents.")
    add_bullets(doc, [
        "Livrables natifs générés par l'application : planning détaillé, WBS, matrice RACI, schéma de communication, budget prévisionnel, PV de kick-off. Ces fichiers sont générés avec le logo CIT.",
        "Pièces externes ou probantes : charte signée, documents tiers, annexes fournisseur, supports non structurés. Ces pièces sont déposées telles quelles.",
        "La présence d'un document à l'écran dépend du type de livrable enregistré lors du dépôt ou de la génération.",
    ])

    add_section_intro(doc, "16. Messages d'alerte et interprétation", "Les alertes doivent être lues comme des guides d'action.")
    add_bullets(doc, [
        "Dossier incomplet : la charte signée, ses signatures ou un prérequis de validation manque.",
        "Ajoutez d'abord au moins une tâche : le planning interactif n'a pas encore de ligne sauvegardée.",
        "Livrables obligatoires manquants : l'écran cite explicitement les livrables ou validations à compléter avant passage de phase.",
        "Validation DSI en attente DM : la validation DSI est verrouillée tant que le Directeur Métier n'a pas validé.",
    ])

    add_section_intro(doc, "17. Recommandations d'usage", "Pour un usage fluide et fiable, respecter l'ordre d'exécution suivant.")
    add_bullets(doc, [
        "Saisir d'abord dans l'application, puis générer les livrables natifs ; éviter la double tenue entre fichiers externes et écrans.",
        "Ne déposer manuellement que les pièces externes, signées ou probantes.",
        "Lire les alertes avant de forcer une validation : elles indiquent exactement ce qui manque.",
        "Toujours enregistrer un écran structuré avant de lancer la génération du document associé.",
    ])

    return doc


def save(doc: Document, target: Path) -> None:
    doc.save(str(target))


def main() -> None:
    save(workflow_doc(), WORKFLOW_DOC)
    save(full_doc(), FULL_DOC)
    print(f"Generated: {WORKFLOW_DOC}")
    print(f"Generated: {FULL_DOC}")


if __name__ == "__main__":
    main()
