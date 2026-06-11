from __future__ import annotations

from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(r"C:\Users\kerne\Downloads\aldric\projets\GestionProjetIT")
TARGET = ROOT / "Documentation_Complete_Application_GestionProjetsIT.docx"

ACCENT = RGBColor(0x17, 0x3D, 0x6B)
ACCENT_2 = RGBColor(0x2F, 0x6F, 0xAD)
TEXT_MUTED = RGBColor(0x6B, 0x72, 0x80)
HEADER_FILL = "EAF2F8"
CALLOUT_FILL = "F5F9FC"
WARN_FILL = "FFF4CC"
SUCCESS_FILL = "E8F5E9"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.12
    normal.paragraph_format.space_after = Pt(6)

    for name, size, color in [
        ("Title", 24, ACCENT),
        ("Heading 1", 16, ACCENT),
        ("Heading 2", 13, ACCENT),
        ("Heading 3", 11, ACCENT_2),
    ]:
        style = doc.styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("Documentation complète - Gestion Projets IT")
    run.font.name = "Arial"
    run.font.size = Pt(8.5)
    run.font.color.rgb = TEXT_MUTED

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("CIT - DSI - 21/05/2026")
    run.font.name = "Arial"
    run.font.size = Pt(8.5)
    run.font.color.rgb = TEXT_MUTED


def add_para(doc: Document, text: str, bold_prefix: str | None = None) -> None:
    p = doc.add_paragraph()
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
    p.add_run(text)


def add_bullet(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="List Bullet")


def add_number(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="List Number")


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths_cm: list[float] | None = None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, text in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = text
        set_cell_shading(cell, HEADER_FILL)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.name = "Arial"
                run.font.size = Pt(9.5)
    for values in rows:
        row = table.add_row()
        for i, text in enumerate(values):
            cell = row.cells[i]
            cell.text = text
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(9.5)
    if widths_cm:
        for row in table.rows:
            for i, width in enumerate(widths_cm):
                row.cells[i].width = Cm(width)
    return table


def add_callout(doc: Document, text: str, fill: str = CALLOUT_FILL) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    cell.text = text
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=110, bottom=110, start=140, end=140)


def cover_page(doc: Document) -> None:
    title = doc.add_paragraph(style="Title")
    title.add_run("Documentation complète de l'application")
    subtitle = doc.add_paragraph()
    r = subtitle.add_run("Gestion Projets IT - Côte d'Ivoire Terminal")
    r.font.name = "Arial"
    r.font.size = Pt(14)
    r.bold = True
    r.font.color.rgb = ACCENT_2

    desc = doc.add_paragraph()
    r = desc.add_run(
        "Document de référence fonctionnel, technique et d'exploitation construit à partir du code réel, "
        "des écrans existants, des workflows actifs, des règles de blocage et de l'architecture actuellement déployée."
    )
    r.font.name = "Arial"
    r.font.size = Pt(11)
    r.font.color.rgb = TEXT_MUTED

    add_table(
        doc,
        ["Champ", "Valeur"],
        [
            ["Organisation", "Côte d'Ivoire Terminal - DSI"],
            ["Application", "Gestion Projets IT"],
            ["Technologie", "ASP.NET Core MVC .NET 9 / SQL Server / Razor / EF Core"],
            ["Date de mise à jour", "21/05/2026"],
            ["Base documentaire", "Code source, vues, contrôleurs, services, tests et configurations du dépôt courant"],
            ["Statut qualité connu", "Build valide et 186 tests automatisés réussis dans la session de travail"],
        ],
        [4.6, 11.6],
    )
    doc.add_paragraph()
    add_callout(
        doc,
        "Usage recommandé : ce document doit servir à la fois de manuel de compréhension globale, "
        "de guide d'exploitation DSI, de support de formation des rôles métiers et de référence de diagnostic "
        "pour le support applicatif.",
        SUCCESS_FILL,
    )
    doc.add_page_break()


def build_doc() -> Document:
    doc = Document()
    configure_document(doc)
    cover_page(doc)

    doc.add_heading("1. Résumé exécutif", level=1)
    add_para(
        doc,
        "Gestion Projets IT est une application web interne destinée à centraliser la gestion des demandes de projet IT, "
        "le suivi du portefeuille, les validations séquentielles métier/DSI, la gestion des phases projet, les livrables, "
        "les risques, les charges, les notifications et la clôture."
    )
    add_table(
        doc,
        ["Indicateur", "Valeur actuelle"],
        [
            ["Framework", ".NET 9 / ASP.NET Core MVC"],
            ["Base de données", "SQL Server via Entity Framework Core 9"],
            ["Authentification", "Cookie local + Azure AD (configurable)"],
            ["Contrôleurs", "11 contrôleurs principaux"],
            ["Vues Razor", "85 vues .cshtml"],
            ["Services métier", "17 services d'infrastructure actifs"],
            ["Entités métier", "35 entités persistées environ"],
            ["Tests automatisés", "186 tests réussis / 186"],
        ],
        [4.5, 11.7],
    )
    add_callout(
        doc,
        "Point de vérité : les workflows décrits ci-après reflètent le comportement actuellement câblé dans les contrôleurs "
        "et les règles de validation serveur, et non uniquement le cahier des charges initial.",
    )

    doc.add_heading("2. Objectifs de l'application", level=1)
    for text in [
        "Formaliser les demandes de projet IT dans un circuit unique.",
        "Assurer la validation séquentielle Directeur Métier puis DSI.",
        "Créer automatiquement les projets après validation DSI.",
        "Piloter les phases Analyse, Planification, Exécution, UAT/MEP et Clôture.",
        "Structurer les livrables obligatoires et bloquer les transitions si le dossier est incomplet.",
        "Donner une visibilité adaptée par rôle, direction et portefeuille.",
        "Tracer les actions sensibles par audit et notifications.",
    ]:
        add_bullet(doc, text)

    doc.add_heading("3. Architecture technique", level=1)
    add_table(
        doc,
        ["Composant", "Description"],
        [
            ["Frontend", "Razor Views + CSS applicatif personnalisé + composants partiels et ViewComponents."],
            ["Backend Web", "ASP.NET Core MVC avec contrôleurs sécurisés et filtres globaux."],
            ["Persistance", "ApplicationDbContext EF Core avec SQL Server."],
            ["Authentification", "Cookie local + OpenID Connect Azure AD."],
            ["Fichiers", "Stockage géré via FileStorageService et accès sécurisé via DocumentController."],
            ["Exports", "Services PDF, Word et Excel dédiés."],
            ["Notifications", "Notifications applicatives internes + services Email et Teams."],
            ["Journalisation", "Serilog + AuditLog métier."],
            ["Sécurité", "Headers HTTP durcis, antiforgery, rate limiting, middleware d'exception."],
        ],
        [4.4, 11.8],
    )

    doc.add_heading("4. Stack logicielle et dépendances", level=1)
    add_table(
        doc,
        ["Dépendance", "Usage"],
        [
            ["Microsoft.EntityFrameworkCore.SqlServer 9.0.11", "ORM et accès SQL Server."],
            ["DocumentFormat.OpenXml 3.0.1", "Manipulation de documents Office."],
            ["EPPlus 7.5.2", "Exports Excel."],
            ["QuestPDF 2025.7.4", "Génération PDF."],
            ["Serilog", "Logs applicatifs fichier + console."],
            ["Microsoft.Identity.Web / OIDC", "Intégration Azure AD."],
            ["BCrypt.Net-Next", "Hash des mots de passe locaux."],
            ["FluentValidation", "Validation métier complémentaire."],
        ],
        [5.0, 11.2],
    )

    doc.add_heading("5. Modèle de données principal", level=1)
    add_para(doc, "Les entités les plus structurantes de l'application sont les suivantes :")
    add_table(
        doc,
        ["Entité", "Rôle fonctionnel"],
        [
            ["Utilisateur / UtilisateurRole", "Comptes, rôles multiples et rattachement organisationnel."],
            ["Direction / Service", "Référentiels métier et rattachement des utilisateurs."],
            ["DemandeProjet", "Demande initiale portée par le demandeur."],
            ["DocumentJointDemande", "Pièces jointes de la demande."],
            ["Projet", "Entité centrale de pilotage après validation DSI."],
            ["FicheProjet", "Données complémentaires par phase (analyse, planification, exécution, clôture)."],
            ["CharteProjet / JalonsCharte / PartiesPrenantesCharte", "Charte du projet et ses sous-composants."],
            ["LivrableProjet", "Documents produits par phase."],
            ["MembreProjet", "Composition de l'équipe projet."],
            ["RisqueProjet / AnomalieProjet", "Registre des risques et anomalies."],
            ["ChargeProjet", "Charges prévues et réelles par période et ressource."],
            ["DemandeClotureProjet", "Circuit de clôture à validations séquentielles."],
            ["HistoriquePhaseProjet / HistoriqueChefProjet", "Traçabilité des changements de phase et d'affectation."],
            ["Notification / AuditLog", "Journal de notification et audit technique/métier."],
            ["RolePermission", "Matrice des autorisations par rôle."],
            ["DelegationValidationDSI / DelegationChefProjet", "Gestion des délégations temporaires."],
            ["DossierSignatureProjet / SignataireDossierSignatureProjet", "Pilotage de signature électronique des chartes."],
            ["CasTestProjet / CampagneTestProjet / ExecutionTestProjet", "Structuration des campagnes de tests UAT."],
            ["CollaborationProjet / TacheCollaborationProjet", "Suivi collaboratif et actions transverses."],
        ],
        [4.8, 11.4],
    )

    doc.add_heading("6. Services métier et services transverses", level=1)
    add_table(
        doc,
        ["Service", "Responsabilité"],
        [
            ["CurrentUserService", "Expose le contexte utilisateur courant."],
            ["PermissionService", "Évalue les permissions écran/action par rôle."],
            ["LivrableValidationService", "Détermine les livrables obligatoires et blocages de passage de phase."],
            ["RAGCalculationService", "Calcule l'état RAG des projets."],
            ["NotificationService", "Crée et diffuse les notifications applicatives."],
            ["TeamsNotificationService", "Prépare les notifications Teams."],
            ["EmailService", "Gestion des emails sortants si activés."],
            ["FileStorageService", "Validation, stockage et récupération sécurisée des fichiers."],
            ["DocumentPreviewService", "Prévisualisation sécurisée des livrables et documents."],
            ["PdfService / WordService / ExcelService", "Génération des exports et documents bureautiques."],
            ["AuditService", "Journal d'audit métier et technique."],
            ["ElectronicSignatureService", "Gestion du dossier de signature électronique."],
            ["UatValidationService", "Vérifications de cohérence en phase UAT."],
            ["CollaborationProjetService", "Fonctions liées à la collaboration projet."],
            ["CacheService", "Abstraction de cache mémoire."],
        ],
        [4.7, 11.5],
    )

    doc.add_heading("7. Sécurité, authentification et autorisations", level=1)
    add_para(
        doc,
        "L'application impose une authentification globale. Les rôles standards sont Demandeur, Directeur Métier, DSI, "
        "Responsable Solutions IT, Chef de Projet et AdminIT."
    )
    add_table(
        doc,
        ["Mécanisme", "Description opérationnelle"],
        [
            ["Cookie auth", "Authentification locale avec expiration glissante, 30 minutes d'inactivité."],
            ["Azure AD", "OpenID Connect configurable via la section AzureAd de appsettings.json."],
            ["Antiforgery", "AutoValidateAntiforgeryToken global sur les contrôleurs."],
            ["Rate limiting", "100 requêtes/minute/IP, 5 tentatives de login/15 minutes, 20 uploads/minute."],
            ["Permission matrix", "Filtre global PermissionMatrixAuthorizationFilter + menu piloté par RolePermission."],
            ["Admin override", "AdminIT reste superutilisateur via handler dédié."],
            ["Security headers", "CSP, nosniff, frame deny, referrer-policy, permissions-policy."],
            ["Audit", "Journalisation des validations, rejets, sauvegardes et actions sensibles."],
        ],
        [4.4, 11.8],
    )
    add_callout(
        doc,
        "Configuration Azure AD actuelle : le code supporte Azure AD, mais l'environnement par défaut contient encore "
        "des placeholders TenantId / ClientId / ClientSecret ; ces valeurs doivent être remplacées en production.",
        WARN_FILL,
    )

    doc.add_heading("8. Rôles applicatifs et périmètres", level=1)
    add_table(
        doc,
        ["Rôle", "Périmètre principal", "Validation autorisée"],
        [
            ["Demandeur", "Crée la demande, suit ses projets, valide la clôture demandeur.", "Clôture demandeur."],
            ["Directeur Métier", "Valide les demandes de sa direction et agit comme sponsor du projet.", "Demande initiale, charte, planification, clôture DM."],
            ["DSI", "Voit tout le portefeuille, valide les demandes, charte, planification et clôture finale.", "Demande initiale, charte, planification, clôture DSI."],
            ["Responsable Solutions IT", "Support transverse et délégation DSI éventuelle.", "Charte, planification, clôture DSI si délégation active."],
            ["Chef de Projet", "Prépare toutes les phases et soumet les validations.", "Passage Analyse -> Planification, soumission clôture."],
            ["AdminIT", "Administration complète et supervision totale.", "Peut tout faire côté application."],
        ],
        [3.5, 7.2, 5.5],
    )

    doc.add_heading("9. Cartographie des modules", level=1)
    add_table(
        doc,
        ["Module", "Description"],
        [
            ["Authentification / Compte", "Login local, Azure AD, demande d'accès, profil, changement de mot de passe."],
            ["Administration", "Utilisateurs, rôles, directions, services, délégations, paramètres, demandes de compte."],
            ["Demandes projet", "Création, brouillons, pièces jointes, validation DM, validation DSI, corrections, historique."],
            ["Portefeuille / tableaux de bord", "Portefeuille projet, dashboard DSI/Admin, analytics et vues direction."],
            ["Projet", "Détails, synthèse, analyse, charte, planification, exécution, UAT, clôture, historique."],
            ["Charges & capacité", "Suivi des charges, capacité par ressource, synthèse prévue/réelle."],
            ["Notifications", "Centre de notifications avec marquage unitaire et global."],
            ["Autorisations & droits", "Matrice des écrans/actions par rôle."],
            ["Aide", "Centre d'aide et guides par rôle."],
        ],
        [4.8, 11.4],
    )

    doc.add_heading("10. Workflow fonctionnel global", level=1)
    add_number(doc, "Le demandeur crée une demande projet et joint le cahier des charges.")
    add_number(doc, "Le Directeur Métier valide, corrige ou rejette la demande.")
    add_number(doc, "La DSI valide la demande et le système crée automatiquement le projet.")
    add_number(doc, "Le Chef de Projet prépare l'analyse, la charte, les signatures et déclenche la validation charte.")
    add_number(doc, "Le Directeur Métier valide la charte, puis la DSI ou le RSIT délégué.")
    add_number(doc, "Le Chef de Projet valide la phase Analyse et le projet entre en Planification & Validation.")
    add_number(doc, "Le Chef de Projet prépare la planification et charge les livrables obligatoires.")
    add_number(doc, "Le Directeur Métier valide la planification, puis la DSI ou le RSIT délégué passent le projet en exécution.")
    add_number(doc, "Le Chef de Projet pilote ensuite Exécution, UAT & MEP, puis soumet la clôture.")
    add_number(doc, "Le demandeur, le Directeur Métier puis la DSI valident la clôture finale.")

    doc.add_heading("11. Phase 1 - Demande initiale", level=1)
    add_para(doc, "La demande initiale couvre le recueil du besoin et le circuit de validation avant création du projet.")
    add_table(
        doc,
        ["Étape", "Écran", "Acteur", "Sortie"],
        [
            ["Création", "DemandeProjet/Create", "Demandeur", "Demande enregistrée avec pièces jointes."],
            ["Validation métier", "DemandeProjet/ListeValidationDM", "Directeur Métier", "Validation, retour correction ou rejet."],
            ["Validation DSI", "DemandeProjet/ListeValidationDSI", "DSI", "Validation finale et création projet."],
            ["Consultation", "DemandeProjet/Index", "Demandeur / DSI / DM", "Suivi de statut et historique."],
        ],
        [2.0, 5.0, 3.2, 5.5],
    )

    doc.add_heading("12. Phase 2 - Analyse & Clarification", level=1)
    add_para(
        doc,
        "L'onglet Analyse est la zone de cadrage. Le CP y dépose les documents d'analyse, gère les risques et prépare la charte."
    )
    add_table(
        doc,
        ["Élément", "Où il se gère", "Valeur métier"],
        [
            ["Cahier d'analyse technique", "Détails projet > Analyse", "Document obligatoire de cadrage technique."],
            ["Note de cadrage", "Détails projet > Analyse", "Document obligatoire de cadrage projet."],
            ["Registre des risques", "Détails projet > Analyse", "Risques saisis et mis à jour par le CP."],
            ["Charte projet", "Projet/CharteProjet/{id}", "Document central de validation de phase Analyse."],
            ["Version signée", "Projet/CharteProjet/{id}", "Condition de complétude du dossier de charte."],
            ["Validation charte", "Projet/ValidationsProjet", "Validation DM puis DSI/RSIT délégué."],
        ],
        [4.2, 5.0, 5.7],
    )
    add_callout(
        doc,
        "Important : DM et DSI ne valident pas directement le Cahier d'analyse technique ou la Note de cadrage. "
        "Ils valident le dossier de charte via l'écran ValidationsProjet, après dépôt d'une charte signée complète.",
        CALLOUT_FILL,
    )
    add_para(doc, "Préconditions de validation Analyse :")
    add_bullet(doc, "Le livrable Charte projet doit exister.")
    add_bullet(doc, "Le livrable Charte projet signée doit exister.")
    add_bullet(doc, "SignatureSponsor et SignatureChefProjet doivent être enregistrées.")
    add_bullet(doc, "Le DM valide d'abord la charte, puis la DSI ou le RSIT délégué.")

    doc.add_heading("13. Phase 3 - Planification & Validation", level=1)
    add_para(
        doc,
        "La planification se gère dans l'onglet Planification du projet. Le CP renseigne les informations de planning, WBS, "
        "RACI, communication et budget, puis charge les livrables obligatoires."
    )
    add_table(
        doc,
        ["Livrable obligatoire", "Exemple"],
        [
            ["Planning détaillé", "Gantt, planning Excel ou export MS Project."],
            ["WBS", "Découpage détaillé des tâches."],
            ["Matrice RACI", "Répartition des responsabilités."],
            ["Schéma de communication", "Fréquence, canaux, publics et gouvernance."],
            ["Budget prévisionnel", "Estimation financière du projet."],
            ["PV de kick-off", "Compte-rendu officiel de lancement."],
        ],
        [5.1, 10.9],
    )
    add_para(doc, "Séquence de validation :")
    add_number(doc, "Le CP complète et enregistre la planification.")
    add_number(doc, "Le CP dépose les 6 livrables obligatoires.")
    add_number(doc, "Le DM clique sur Valider la planification.")
    add_number(doc, "La DSI ou le RSIT délégué clique sur Passer en exécution.")
    add_number(doc, "Le projet passe en phase Exécution & Suivi.")

    doc.add_heading("14. Phase 4 - Exécution & Suivi", level=1)
    add_para(
        doc,
        "En exécution, le CP suit l'avancement, les risques, les comptes-rendus, les anomalies et les charges. "
        "Le RAG est recalculé et sert d'indicateur de santé."
    )
    add_table(
        doc,
        ["Bloc", "Contenu"],
        [
            ["Avancement", "Pourcentage d'avancement, commentaire et justification des retards."],
            ["RAG", "État vert / orange / rouge calculé et ajusté."],
            ["Risques", "Ajout, mise à jour, plan de mitigation, responsable."],
            ["Livrables", "Comptes-rendus, rapports et pièces d'exécution."],
            ["Charges", "Prévu, réel, capacité, disponibilité, commentaires."],
        ],
        [4.0, 12.0],
    )
    add_callout(doc, "Le passage Exécution -> UAT est actuellement bloqué au minimum par la présence d'un Compte-rendu de réunion.")

    doc.add_heading("15. Phase 5 - UAT & MEP", level=1)
    add_para(
        doc,
        "La phase UAT & MEP structure la recette, les anomalies, la mise en production et l'hypercare. "
        "Le projet ne peut passer en clôture que si les livrables de recette et de MEP requis sont présents."
    )
    add_table(
        doc,
        ["Livrable requis", "Finalité"],
        [
            ["Cahier de tests", "Base des cas de tests et scénarios."],
            ["Feuille d'anomalies", "Suivi des anomalies et corrections."],
            ["PV de recette", "Validation de recette."],
            ["Dossier MEP", "Plan de mise en production."],
            ["PV MEP", "Constat de mise en production."],
            ["Rapport hypercare", "Suivi post-MEP et stabilisation."],
        ],
        [5.1, 10.9],
    )

    doc.add_heading("16. Phase 6 - Clôture & Leçons apprises", level=1)
    add_para(
        doc,
        "La clôture se gère dans l'onglet Clôture. Le CP renseigne le bilan, le statut final, les leçons apprises "
        "et soumet une demande de clôture."
    )
    add_table(
        doc,
        ["Niveau", "Écran", "Règle actuelle"],
        [
            ["1", "ListeValidationClotureDemandeur", "Le demandeur doit valider en premier."],
            ["2", "ListeValidationClotureDM", "Le Directeur Métier valide après le demandeur."],
            ["3", "ListeValidationClotureDSI", "La DSI ou le RSIT délégué valident en dernier."],
        ],
        [1.2, 6.0, 8.8],
    )
    add_bullet(doc, "La clôture finale ne se fait pas sur un simple changement de statut manuel.")
    add_bullet(doc, "Le workflow complet Demandeur -> DM -> DSI est maintenant respecté côté code.")

    doc.add_heading("17. Livrables par phase", level=1)
    add_table(
        doc,
        ["Phase", "Livrables structurants"],
        [
            ["Analyse & Clarification", "Cahier d'analyse technique, Note de cadrage, Charte projet, Charte projet signée."],
            ["Planification & Validation", "Planning détaillé, WBS, Matrice RACI, Schéma de communication, Budget prévisionnel, PV de kick-off."],
            ["Exécution & Suivi", "Comptes-rendus de réunion, rapports d'avancement, documents de pilotage."],
            ["UAT & MEP", "Cahier de tests, feuille d'anomalies, PV de recette, dossier MEP, PV MEP, rapport hypercare."],
            ["Clôture & Leçons apprises", "Rapport de clôture, PV de clôture, dossier d'exploitation, leçons apprises."],
        ],
        [3.9, 12.1],
    )

    doc.add_heading("18. Dashboards, portefeuille et analytics", level=1)
    add_para(
        doc,
        "L'application propose un cockpit dashboard DSI/Admin et des vues directionnelles avec indicateurs projets, "
        "santé du portefeuille, gouvernance, ressources & capacité, finances et validations en attente."
    )
    add_bullet(doc, "Portefeuille DSI avec filtres par direction, phase, statut, CP et santé.")
    add_bullet(doc, "Dashboard Home avec KPI, tableaux critiques, répartitions et synthèses.")
    add_bullet(doc, "Analytics direction pour lecture sponsor / métier.")
    add_bullet(doc, "Historique & traçabilité avec vues consolidées par direction.")

    doc.add_heading("19. Notifications et audit", level=1)
    add_para(
        doc,
        "Chaque étape sensible peut générer une notification interne. Les actions importantes sont également journalisées dans les audits."
    )
    add_table(
        doc,
        ["Mécanisme", "Comportement"],
        [
            ["Notifications internes", "Créées lors des validations, rejets, demandes de correction, soumissions de clôture."],
            ["Centre de notifications", "Vue dédiée avec marquage unitaire ou global comme lu."],
            ["AuditLog", "Journal des validations, refus, sauvegardes, affectations et modifications sensibles."],
            ["Logs techniques", "Logs Serilog dans logs/gestion-projets-*.txt."],
        ],
        [4.8, 11.4],
    )

    doc.add_heading("20. Autorisations et menu dynamique", level=1)
    add_para(
        doc,
        "Le module Autorisations / Droits pilote désormais réellement le menu et le blocage backend des écrans. "
        "L'AdminIT peut ajuster les accès par rôle dans une matrice de permissions."
    )
    add_bullet(doc, "Le menu n'est plus seulement codé en dur ; il s'aligne sur la matrice de permissions.")
    add_bullet(doc, "Le filtre PermissionMatrixAuthorizationFilter bloque l'accès direct aux URLs non autorisées.")
    add_bullet(doc, "AdminIT conserve un accès total via un handler d'override dédié.")

    doc.add_heading("21. Gestion documentaire et fichiers", level=1)
    add_para(
        doc,
        "Les fichiers sont validés et servis de manière sécurisée. Les aperçus passent par DocumentController "
        "et les téléchargements vérifient l'accès projet."
    )
    add_bullet(doc, "Validation des extensions, taille et signature de fichier côté stockage.")
    add_bullet(doc, "Prévisualisation sécurisée des livrables par projet.")
    add_bullet(doc, "Téléchargement direct et ouvertures PDF inline pour certains exports.")

    doc.add_heading("22. Configuration applicative", level=1)
    add_table(
        doc,
        ["Clé / zone", "Usage"],
        [
            ["ConnectionStrings:DefaultConnection", "Connexion SQL Server principale."],
            ["SmtpSettings:*", "Activation et paramètres du serveur SMTP."],
            ["AzureAd:*", "Configuration OpenID Connect Azure AD."],
            ["AllowedHosts", "Hôtes autorisés par ASP.NET Core."],
            ["ParametreSysteme", "Référentiel en base pour certains paramètres dynamiques."],
        ],
        [5.5, 10.5],
    )
    add_callout(
        doc,
        "Valeurs observées dans appsettings.json : SMTP désactivé par défaut et Azure AD encore configuré avec des placeholders. "
        "Un paramétrage de production reste donc nécessaire.",
        WARN_FILL,
    )

    doc.add_heading("23. Déploiement, démarrage et maintenance", level=1)
    add_number(doc, "Configurer la chaîne de connexion SQL Server.")
    add_number(doc, "Renseigner les paramètres Azure AD et SMTP si utilisés.")
    add_number(doc, "Démarrer l'application ASP.NET Core ; les migrations EF et certains patches SQL sont exécutés au démarrage.")
    add_number(doc, "Vérifier les dossiers de logs et le stockage documentaire.")
    add_number(doc, "Vérifier les droits par rôle dans Autorisations / Droits et, si nécessaire, utiliser Réinitialiser.")

    doc.add_heading("24. Qualité logicielle et tests", level=1)
    add_para(
        doc,
        "L'état de qualité vérifié dans la session courante indique que les tests unitaires et d'intégration disponibles passent intégralement."
    )
    add_table(
        doc,
        ["Commande", "Résultat observé"],
        [
            ["dotnet test Tests/GestionProjects.Tests.csproj --no-build --no-restore", "186 réussites, 0 échec, 0 ignoré."],
        ],
        [7.8, 8.2],
    )
    add_bullet(doc, "Le dépôt contient également plusieurs rapports markdown et CSV de vérification et d'implémentation.")
    add_bullet(doc, "Les tests couvrent l'authentification, la sécurité, les demandes, les validations, les phases projet et les services métier.")

    doc.add_heading("25. Guide de prise en main par acteur", level=1)
    add_table(
        doc,
        ["Acteur", "Démarrage recommandé"],
        [
            ["Demandeur", "Créer une demande, suivre son traitement, puis surveiller ses validations de clôture."],
            ["Directeur Métier", "Traiter Validation DM, consulter Projets de ma direction, valider charte et planification du projet sponsorisé."],
            ["DSI", "Traiter Validation DSI, suivre le portefeuille et les validations finales."],
            ["RSIT", "Consulter les projets, agir techniquement et valider en délégation si activé."],
            ["Chef de Projet", "Travailler dans les onglets de phase du projet et préparer tous les livrables/validations."],
            ["AdminIT", "Paramétrer les référentiels, comptes, droits et délégations."],
        ],
        [3.8, 12.2],
    )

    doc.add_heading("26. Bonnes pratiques d'exploitation", level=1)
    for text in [
        "Toujours déposer les livrables avec le bon TypeLivrable ; l'affichage par phase dépend de ce typage.",
        "Faire valider la charte uniquement après dépôt de la version signée complète.",
        "Utiliser la checklist planification avant de solliciter le DM puis la DSI.",
        "Vérifier les notifications et les commentaires de refus avant de relancer une validation.",
        "Conserver la matrice des droits alignée sur l'organisation réelle et les délégations actives.",
        "Surveiller les logs applicatifs et l'audit lors des incidents ou anomalies de workflow.",
    ]:
        add_bullet(doc, text)

    doc.add_heading("27. Limites et points d'attention", level=1)
    add_bullet(doc, "L'activation réelle d'Azure AD dépend encore du paramétrage de production.")
    add_bullet(doc, "Le typage de certains livrables repose sur le choix utilisateur à l'upload ; le contenu métier du fichier n'est pas inféré automatiquement.")
    add_bullet(doc, "Certains patches de schéma SQL complémentaires sont encore appliqués au démarrage pour sécuriser les environnements existants.")
    add_bullet(doc, "Le moteur de rendu DOCX n'était pas disponible dans cette session ; la présente documentation a donc été vérifiée structurellement mais pas rendue visuellement page par page.")

    doc.add_heading("28. Glossaire", level=1)
    add_table(
        doc,
        ["Terme", "Définition"],
        [
            ["DM", "Directeur Métier, sponsor du projet."],
            ["DSI", "Direction des Systèmes d'Information, valideur transverse et pilote du portefeuille."],
            ["RSIT", "Responsable Solutions IT, valideur DSI par délégation si activé."],
            ["CP", "Chef de Projet DSI."],
            ["RAG", "Indicateur rouge / orange / vert de santé du projet."],
            ["UAT", "User Acceptance Testing / recette."],
            ["MEP", "Mise en production."],
            ["WBS", "Work Breakdown Structure / découpage structuré des tâches."],
            ["RACI", "Répartition Responsable / Approbateur / Consulté / Informé."],
        ],
        [3.5, 12.5],
    )

    doc.add_heading("29. Conclusion", level=1)
    add_para(
        doc,
        "L'application Gestion Projets IT constitue désormais un socle unifié pour piloter les demandes, les validations, "
        "les phases projet, les charges, les livrables, les audits et la clôture. La présente documentation est conçue pour "
        "servir de base commune entre la DSI, les chefs de projet, les directions métiers et l'administration applicative."
    )
    return doc


def main() -> None:
    doc = build_doc()
    props = doc.core_properties
    props.title = "Documentation complète de l'application Gestion Projets IT"
    props.subject = "Documentation fonctionnelle, technique et d'exploitation"
    props.author = "OpenAI Codex"
    props.comments = "Document généré à partir du code réel du dépôt au 21/05/2026."
    props.created = datetime(2026, 5, 21, 10, 0, 0)
    props.modified = datetime(2026, 5, 21, 10, 0, 0)
    doc.save(TARGET)


if __name__ == "__main__":
    main()
