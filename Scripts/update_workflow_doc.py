from __future__ import annotations

import shutil
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(r"C:\Users\kerne\Downloads\aldric\projets\GestionProjetIT")
TARGET = ROOT / "Workflow_Par_Role_GestionProjetsIT.docx"
BACKUP = ROOT / "Workflow_Par_Role_GestionProjetsIT.backup-20260520.docx"

ACCENT = RGBColor(0x1F, 0x4E, 0x79)
ACCENT_LIGHT = "D9EAF7"
HEADER_FILL = "EAF2F8"
WARN_FILL = "FFF4CC"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_document_defaults(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.font.size = Pt(10.5)
    pf = normal.paragraph_format
    pf.space_after = Pt(6)
    pf.line_spacing = 1.12

    for style_name, size, bold, color in [
        ("Title", 24, True, ACCENT),
        ("Heading 1", 16, True, ACCENT),
        ("Heading 2", 13, True, ACCENT),
        ("Heading 3", 11, True, ACCENT),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = bold
        style.font.color.rgb = color


def add_header_footer(section) -> None:
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("Workflow par rôle - Gestion Projets IT")
    run.font.name = "Arial"
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("Côte d'Ivoire Terminal - DSI - Mise à jour 20/05/2026")
    run.font.name = "Arial"
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)


def add_title_page(doc: Document) -> None:
    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.add_run("Workflow par rôle - Gestion Projets IT")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = subtitle.add_run(
        "Version actualisée sur la base du comportement réellement implémenté "
        "dans l'application Gestion Projets IT."
    )
    r.font.name = "Arial"
    r.font.size = Pt(11.5)
    r.font.color.rgb = RGBColor(0x4B, 0x55, 0x63)

    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = note.add_run(
        "Référence de vérité : contrôleurs, vues Razor, règles de blocage et "
        "droits actifs observés dans l'application au 20/05/2026."
    )
    run.italic = True
    run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    table = doc.add_table(rows=4, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    rows = [
        ("Périmètre", "Cycle complet : demande, analyse, planification, exécution, UAT/MEP, clôture."),
        ("Application", "GestionProjetIT - ASP.NET Core MVC / SQL Server."),
        ("Version du document", "Révision actualisée - 20/05/2026."),
        ("Objet", "Décrire, par rôle, les écrans, actions, prérequis et validations réellement en place."),
    ]
    for row, (left, right) in zip(table.rows, rows):
        row.cells[0].text = left
        row.cells[1].text = right
        row.cells[0].width = Cm(4.5)
        row.cells[1].width = Cm(11.8)
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)
        set_cell_shading(row.cells[0], HEADER_FILL)

    doc.add_paragraph()
    warning = doc.add_table(rows=1, cols=1)
    warning.style = "Table Grid"
    cell = warning.cell(0, 0)
    cell.text = (
        "Important : en cas d'écart entre ce document et l'application, le code "
        "et les contrôles serveur prévalent. Le présent document a précisément été "
        "réécrit pour réduire ces écarts."
    )
    set_cell_shading(cell, WARN_FILL)
    set_cell_margins(cell, top=110, bottom=110, start=140, end=140)

    doc.add_page_break()


def add_heading(doc: Document, text: str, level: int) -> None:
    doc.add_heading(text, level=level)


def add_body(doc: Document, text: str, bold_prefix: str | None = None) -> None:
    p = doc.add_paragraph()
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
    p.add_run(text)


def add_bullet(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="List Bullet")


def add_number(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="List Number")


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths_cm: list[float] | None = None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    hdr = table.rows[0]
    for idx, text in enumerate(headers):
        cell = hdr.cells[idx]
        cell.text = text
        set_cell_shading(cell, HEADER_FILL)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.name = "Arial"
                run.font.size = Pt(9.5)
    set_repeat_table_header(hdr)

    for values in rows:
        row = table.add_row()
        for idx, text in enumerate(values):
            cell = row.cells[idx]
            cell.text = text
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(9.5)

    if widths_cm:
        for row in table.rows:
            for idx, width in enumerate(widths_cm):
                row.cells[idx].width = Cm(width)

    return table


def add_role_section(doc: Document, title: str, mission: str, screens: list[list[str]], steps: list[list[str]]) -> None:
    add_heading(doc, title, 2)
    add_body(doc, mission)
    add_heading(doc, "Écrans principaux", 3)
    add_table(doc, ["Écran", "Utilité", "Observation"], screens, [4.3, 7.2, 5.0])
    add_heading(doc, "Workflow opérationnel", 3)
    add_table(doc, ["Étape", "Écran", "Action", "Sortie / condition"], steps, [1.2, 4.0, 5.8, 5.5])


def build_document() -> Document:
    doc = Document()
    set_document_defaults(doc)
    add_header_footer(doc.sections[0])
    add_title_page(doc)

    add_heading(doc, "1. Objet et conventions", 1)
    add_body(
        doc,
        "Ce document décrit le workflow réellement utilisé dans l'application Gestion Projets IT, "
        "en se basant sur les écrans disponibles, les validations serveur, les routes actives et "
        "les blocages automatiques en place."
    )
    add_bullet(doc, "DM = Directeur Métier et sponsor du projet.")
    add_bullet(doc, "DSI = valideur transverse et propriétaire du portefeuille.")
    add_bullet(doc, "RSIT = Responsable Solutions IT ; il valide comme la DSI uniquement si une délégation est active.")
    add_bullet(doc, "CP = Chef de Projet DSI assigné au projet.")
    add_bullet(doc, "AdminIT = superutilisateur applicatif et administrateur des paramètres.")

    add_heading(doc, "2. Vue d'ensemble du cycle global", 1)
    add_body(
        doc,
        "Le cycle métier complet se déroule en six temps : demande initiale, analyse, planification, "
        "exécution, UAT/MEP et clôture. À partir du moment où la demande DSI est validée, le projet suit "
        "les phases d'exécution gérées dans le module Projet."
    )
    add_table(
        doc,
        ["Étape", "Nom métier", "Acteurs dominants", "Sortie attendue"],
        [
            ["1", "Soumission et validation de la demande", "Demandeur, DM, DSI", "Projet créé après validation DSI."],
            ["2", "Analyse & Clarification", "CP, DM, DSI/RSIT", "Charte préparée, signée, validée DM puis DSI."],
            ["3", "Planification & Validation", "CP, DM, DSI/RSIT", "Dossier de planification complet et validé."],
            ["4", "Exécution & Suivi", "CP, DSI/RSIT", "Pilotage opérationnel, risques, charges, comptes-rendus."],
            ["5", "UAT & MEP", "CP, métiers, DSI/RSIT", "Recette, MEP et hypercare documentés."],
            ["6", "Clôture & Leçons apprises", "CP, Demandeur, DM, DSI", "Projet clôturé après triple validation finale."],
        ],
        [1.0, 5.2, 5.1, 5.2],
    )

    add_heading(doc, "3. Matrice synthétique des rôles", 1)
    add_table(
        doc,
        ["Rôle", "Mission principale", "Peut valider", "Écrans pivots"],
        [
            ["Demandeur", "Créer la demande et valider la clôture côté métier demandeur.", "Clôture demandeur", "Mes Demandes, Mes validations clôture"],
            ["Directeur Métier", "Valider la demande initiale et agir comme sponsor.", "Demande initiale, charte, planification, clôture DM", "Validation DM, Validations Projet, Planification, Clôture DM"],
            ["DSI", "Arbitrer les demandes et piloter les validations transverses.", "Demande initiale, charte, planification, clôture DSI", "Validation DSI, Validations Projet, Planification, Clôture DSI"],
            ["RSIT", "Support technique transverse et délégation DSI éventuelle.", "Charte, planification, clôture DSI si délégation active", "Validations Projet, Portefeuille, détails projet"],
            ["Chef de Projet", "Préparer le dossier projet et piloter toutes les phases.", "Passage de phase Analyse, soumission clôture", "Détails projet, CharteProjet, Charges, UAT, Clôture"],
            ["AdminIT", "Administration applicative et supervision complète.", "Tous les circuits si nécessaire", "Administration, Autorisations, Users, paramètres"],
        ],
        [3.0, 5.4, 4.4, 5.0],
    )

    add_heading(doc, "4. Workflow détaillé par rôle", 1)
    add_role_section(
        doc,
        "4.1 Demandeur",
        "Le demandeur initie le besoin, répond aux retours éventuels et intervient à la fin du cycle pour valider la clôture côté demandeur.",
        [
            ["Nouvelle demande", "Créer une demande projet avec pièces jointes.", "Point d'entrée du cycle."],
            ["Mes Demandes", "Suivre l'état de ses demandes et corriger si demandé.", "Le demandeur ne valide pas la charte projet."],
            ["Mes validations clôture", "Traiter les demandes de clôture en attente.", "Première validation de clôture."],
        ],
        [
            ["1", "Nouvelle demande", "Saisir le besoin, joindre le cahier des charges, soumettre.", "La demande part au DM."],
            ["2", "Mes Demandes", "Répondre aux demandes de correction du DM ou de la DSI.", "La demande retourne ensuite au circuit de validation."],
            ["3", "Mes Projets / détails", "Consulter l'avancement du projet après création.", "Vue de suivi, pas de validation de phase."],
            ["4", "ListeValidationClotureDemandeur", "Valider la clôture une fois la demande soumise par le CP.", "La clôture passe ensuite au DM."],
        ],
    )
    add_role_section(
        doc,
        "4.2 Directeur Métier (Sponsor)",
        "Le Directeur Métier agit à la fois comme valideur de la demande initiale et comme sponsor du projet pendant l'analyse, la planification et la clôture.",
        [
            ["Validation DM", "Valider ou corriger la demande initiale.", "Écran distinct du projet."],
            ["Validations Projet", "Valider la charte de phase Analyse.", "Bouton vert ouvert seulement si le dossier de charte est complet."],
            ["Détails projet > Planification", "Valider la planification du projet sponsorisé.", "Visible si le DM est sponsor du projet."],
            ["ListeValidationClotureDM", "Valider ou rejeter la clôture après validation du demandeur.", "Deuxième niveau de clôture."],
        ],
        [
            ["1", "ListeValidationDM", "Valider, corriger ou rejeter la demande initiale.", "En cas de validation, la demande part à la DSI."],
            ["2", "ValidationsProjet", "Consulter le dossier de charte et valider côté DM.", "Nécessite une charte signée complète."],
            ["3", "Details?id&tab=planification", "Cliquer sur Valider la planification.", "Ouvre ensuite la validation DSI."],
            ["4", "ListeValidationClotureDM", "Valider ou refuser la clôture après validation du demandeur.", "En cas de validation, le dossier part à la DSI."],
        ],
    )
    add_role_section(
        doc,
        "4.3 DSI",
        "La DSI valide la demande initiale, la charte, la planification et la clôture finale. Elle voit l'ensemble du portefeuille.",
        [
            ["Validation DSI", "Traitement des demandes initiales après DM.", "Peut déléguer au RSIT."],
            ["Validations Projet", "Validation DSI de la charte de phase Analyse.", "Disponible seulement après validation DM."],
            ["Détails projet > Planification", "Validation DSI de la planification.", "Le projet passe en exécution après succès."],
            ["ListeValidationClotureDSI", "Validation finale de clôture.", "Dernier niveau de clôture."],
        ],
        [
            ["1", "ListeValidationDSI", "Valider ou rejeter la demande après DM.", "En cas de validation, création du projet."],
            ["2", "ValidationsProjet", "Valider la charte une fois le DM passé et le dossier complet.", "La phase Analyse devient clôturable par le CP."],
            ["3", "Details?id&tab=planification", "Cliquer sur Passer en exécution.", "Nécessite validation DM et 6 livrables obligatoires."],
            ["4", "ListeValidationClotureDSI", "Valider ou rejeter la clôture finale.", "Le projet passe au statut Clôturé."],
        ],
    )
    add_role_section(
        doc,
        "4.4 Responsable Solutions IT",
        "Le RSIT consulte transversalement les projets et n'endosse le rôle de valideur DSI que si une délégation DSI active lui a été attribuée.",
        [
            ["Portefeuille / détails", "Consultation transverse et appui technique.", "Toujours disponible."],
            ["Validations Projet", "Valider comme la DSI si délégation active.", "Même écran et mêmes blocages que la DSI."],
            ["Détails projet > Planification", "Valider la planification comme la DSI si délégation active.", "Même règle que la DSI."],
            ["ListeValidationClotureDSI", "Clôture DSI par délégation.", "Accessible uniquement en délégation."],
        ],
        [
            ["1", "Portefeuille / Details", "Analyser les risques, charges et livrables techniques.", "Rôle de support et de supervision."],
            ["2", "ValidationsProjet", "Valider la charte à la place de la DSI si délégation active.", "Le DM doit déjà avoir validé."],
            ["3", "Details?id&tab=planification", "Passer la planification en exécution si délégation active.", "Les livrables de planification sont obligatoires."],
            ["4", "ListeValidationClotureDSI", "Valider la clôture finale comme délégataire DSI.", "Intervient après demandeur et DM."],
        ],
    )
    add_role_section(
        doc,
        "4.5 Chef de Projet",
        "Le Chef de Projet prépare le dossier à chaque phase, dépose les livrables, pilote les risques et les charges, puis soumet les validations métier.",
        [
            ["Détails projet > Analyse", "Zone de travail de phase Analyse.", "Documents d'analyse, risques, déclenchement du passage de phase."],
            ["CharteProjet", "Éditer la charte, générer le PDF, déposer la version signée.", "Le CP gère aussi les cases de signature."],
            ["Détails projet > Planification", "Préparer le dossier de planification et déposer les livrables.", "Les validations DM/DSI sont sur le même écran projet."],
            ["Détails projet > Exécution / Charges", "Suivre avancement, risques, charges et comptes-rendus.", "Pilotage hebdomadaire."],
            ["Détails projet > UAT & MEP", "Conduire recette, MEP et hypercare.", "Prépare le passage en clôture."],
            ["Détails projet > Clôture", "Soumettre la demande de clôture.", "Déclenche le circuit demandeur puis DM puis DSI."],
        ],
        [
            ["1", "Details?id&tab=analyse", "Compléter le cadrage, déposer les livrables d'analyse.", "Le CP prépare, mais ne valide pas seul la charte."],
            ["2", "CharteProjet", "Sauvegarder la charte, générer le PDF, déposer la charte signée, enregistrer les signatures.", "Le dossier de charte devient prêt pour DM/DSI."],
            ["3", "ValidationsProjet", "Suivre le statut de validation charte.", "Le CP n'y valide pas en tant que DM ou DSI."],
            ["4", "ValiderPhaseAnalyse", "Passer en planification après validations DM + DSI et livrables requis.", "Le projet entre en Planification & Validation."],
            ["5", "UpdatePlanification", "Renseigner la planification et déposer les 6 livrables obligatoires.", "Le DM puis la DSI peuvent valider."],
            ["6", "Execution / Charges", "Piloter l'exécution, les risques, les charges et les comptes-rendus.", "Le projet peut ensuite aller en UAT/MEP."],
            ["7", "UAT / Clôture", "Finaliser recette, MEP, hypercare et soumettre la clôture.", "Le circuit final demandeur > DM > DSI s'ouvre."],
        ],
    )
    add_role_section(
        doc,
        "4.6 AdminIT",
        "L'AdminIT est le superutilisateur applicatif. Il peut gérer les utilisateurs, paramètres, délégations, autorisations et intervenir transversalement si nécessaire.",
        [
            ["Administration", "Utilisateurs, rôles, directions, services, délégations, paramètres.", "Administration fonctionnelle et technique."],
            ["Autorisations / Droits", "Matrice des permissions par rôle.", "Le menu et le backend sont désormais pilotés par cette matrice."],
            ["Tous les écrans projet", "Supervision ou reprise de main ponctuelle.", "AdminIT conserve un accès total."],
        ],
        [
            ["1", "Users / Rôles / Directions / Services", "Paramétrer les référentiels et les accès.", "Prépare les conditions d'exploitation."],
            ["2", "Autorisations/Index", "Activer ou couper des écrans par rôle.", "Effet sur le menu et le contrôle d'accès."],
            ["3", "Délégations", "Déléguer la DSI ou le CP selon les besoins.", "Le RSIT peut alors valider comme DSI."],
        ],
    )

    add_heading(doc, "5. Focus phase Analyse & Clarification", 1)
    add_body(
        doc,
        "L'onglet Analyse est une zone de préparation. Les documents qui y sont chargés servent au cadrage, "
        "mais la validation formelle du DM et de la DSI porte sur le dossier de charte, pas sur chaque pièce d'analyse individuellement."
    )
    add_table(
        doc,
        ["Élément", "Où il se gère", "Qui le prépare / dépose", "Rôle dans le workflow"],
        [
            ["Cahier d'analyse technique", "Détails projet > Analyse", "CP / DSI / AdminIT", "Document support de cadrage."],
            ["Note de cadrage", "Détails projet > Analyse", "CP / DSI / AdminIT", "Document support de cadrage."],
            ["Charte projet", "CharteProjet puis Générer PDF Charte", "CP", "Livrable obligatoire pour passer en planification."],
            ["Charte projet signée", "CharteProjet", "CP / DSI / RSIT / AdminIT", "Livrable obligatoire pour la validation DM/DSI."],
            ["Cases Signature sponsor et Signature CP", "CharteProjet > Enregistrer les signatures", "CP / DSI / RSIT / AdminIT", "Condition de complétude du dossier de charte."],
        ],
        [4.5, 4.5, 4.2, 4.0],
    )
    add_body(doc, "Règle de blocage Analyse :")
    add_bullet(doc, "Le passage Analyse -> Planification exige deux livrables obligatoires : Charte projet et Charte projet signée.")
    add_bullet(doc, "La validation DM sur ValidationsProjet exige en plus que la charte signée soit complète : signature Sponsor/DM et signature Chef de Projet.")
    add_bullet(doc, "La validation DSI ou RSIT délégué ne s'ouvre qu'après validation DM.")

    add_heading(doc, "6. Focus phase Planification & Validation", 1)
    add_body(
        doc,
        "La planification se gère dans l'onglet Planification du projet. Le CP y complète la fiche de planification, "
        "dépose les livrables obligatoires puis soumet implicitement le dossier à la validation métier."
    )
    add_table(
        doc,
        ["Bloc", "Données / livrables attendus", "Écran / action", "Observations"],
        [
            ["Planning & jalons", "Prochain jalon, jalons principaux.", "UpdatePlanification", "Renseigné par le CP."],
            ["WBS / ressources / RACI", "Découpage des lots, ressources projet, matrice RACI.", "UpdatePlanification", "Renseigné par le CP."],
            ["Communication & gouvernance", "Fréquence des réunions, participants, canal, COPIL.", "UpdatePlanification", "Renseigné par le CP."],
            ["Budget & risques initiaux", "Budget prévisionnel, commentaire budgétaire, commentaire de validation, synthèse des risques.", "UpdatePlanification", "Renseigné par le CP."],
            ["Livrables obligatoires", "Planning détaillé, WBS, Matrice RACI, Schéma de communication, Budget prévisionnel, PV de kick-off.", "Ajouter un livrable", "Checklist visible dans l'onglet Planification."],
            ["Validation DM", "Clique Valider la planification.", "ValiderPlanifDM", "Réservé au sponsor DM."],
            ["Validation DSI / RSIT délégué", "Clique Passer en exécution.", "ValiderPlanifDSI", "Réservé à la DSI ou au RSIT délégué ; bloqué si un livrable manque."],
        ],
        [3.1, 6.5, 3.2, 4.2],
    )
    add_body(doc, "Règle de blocage Planification :")
    add_bullet(doc, "La DSI ou le RSIT délégué ne peuvent pas valider tant que le DM n'a pas validé.")
    add_bullet(doc, "Le passage Planification -> Exécution est bloqué si l'un des 6 livrables obligatoires manque.")

    add_heading(doc, "7. Focus Exécution, UAT & MEP, Clôture", 1)
    add_table(
        doc,
        ["Phase", "Ce que fait principalement le CP", "Livrable de passage contrôlé", "Validation / sortie"],
        [
            ["Exécution & Suivi", "Mettre à jour avancement, risques, décisions, comptes-rendus, charges.", "Compte-rendu de réunion", "Le projet peut passer en UAT/MEP."],
            ["UAT & MEP", "Piloter recette, anomalies, MEP et hypercare.", "Cahier de tests, Feuille d'anomalies, PV de recette, Dossier MEP, PV MEP, Rapport hypercare", "Le projet peut passer en Clôture & Leçons apprises."],
            ["Clôture & Leçons apprises", "Renseigner bilan, statut final, leçons apprises et soumettre la clôture.", "Rapport de clôture, PV de clôture, dossier d'exploitation selon usage projet", "La demande de clôture part au demandeur, puis DM, puis DSI."],
        ],
        [3.0, 6.8, 4.8, 3.1],
    )
    add_body(doc, "Workflow de clôture réellement implémenté :")
    add_number(doc, "Le CP soumet une demande de clôture depuis l'onglet Clôture.")
    add_number(doc, "Le demandeur valide sur l'écran ListeValidationClotureDemandeur.")
    add_number(doc, "Le DM valide ou rejette sur l'écran ListeValidationClotureDM.")
    add_number(doc, "La DSI ou le RSIT délégué valident ou rejettent sur l'écran ListeValidationClotureDSI.")
    add_number(doc, "En cas de validation finale, le projet passe au statut Clôturé.")

    add_heading(doc, "8. Écrans clés et routes utiles", 1)
    add_table(
        doc,
        ["Écran", "Route type", "Acteurs principaux", "Utilité"],
        [
            ["Nouvelle demande", "/DemandeProjet/Create", "Demandeur", "Créer une demande projet."],
            ["Validation DM", "/DemandeProjet/ListeValidationDM", "DM", "Valider ou corriger la demande initiale."],
            ["Validation DSI", "/DemandeProjet/ListeValidationDSI", "DSI", "Valider ou rejeter la demande après DM."],
            ["Détails projet - Analyse", "/Projet/Details/{id}?tab=analyse", "CP, DM, DSI, RSIT", "Préparer l'analyse et voir les alertes de blocage."],
            ["Charte projet", "/Projet/CharteProjet/{id}", "CP, DSI, RSIT, AdminIT", "Éditer, générer, déposer et signer la charte."],
            ["Validations Projet", "/Projet/ValidationsProjet", "DM, DSI, RSIT délégué", "Valider la charte de phase Analyse."],
            ["Détails projet - Planification", "/Projet/Details/{id}?tab=planification", "CP, DM, DSI, RSIT", "Préparer puis valider la planification."],
            ["Détails projet - Exécution", "/Projet/Details/{id}?tab=execution", "CP, DSI, RSIT", "Suivre avancement, risques et actions."],
            ["Détails projet - UAT", "/Projet/Details/{id}?tab=uat", "CP, métiers, DSI", "Piloter tests, anomalies, recette et MEP."],
            ["Détails projet - Clôture", "/Projet/Details/{id}?tab=cloture", "CP, Demandeur, DM, DSI", "Soumettre et suivre la clôture."],
            ["Clôture demandeur", "/Projet/ListeValidationClotureDemandeur", "Demandeur", "Première validation finale."],
            ["Clôture DM", "/Projet/ListeValidationClotureDM", "DM", "Deuxième validation finale."],
            ["Clôture DSI", "/Projet/ListeValidationClotureDSI", "DSI, RSIT délégué", "Validation finale du projet."],
            ["Autorisations / Droits", "/Autorisations/Index", "AdminIT", "Piloter les droits par rôle."],
        ],
        [3.9, 4.8, 3.6, 4.5],
    )

    add_heading(doc, "9. Règles de blocage à connaître", 1)
    add_table(
        doc,
        ["Point de contrôle", "Blocage automatique", "Comment le lever"],
        [
            ["Analyse -> Planification", "Charte projet ou charte projet signée manquante.", "Générer le PDF charte puis déposer la version signée."],
            ["Validation DM de la charte", "Dossier de charte incomplet.", "Déposer la charte signée et enregistrer les signatures Sponsor/DM et CP."],
            ["Validation DSI de la charte", "Validation DM absente ou charte incomplète.", "Faire valider d'abord le DM puis revenir sur ValidationsProjet."],
            ["Planification -> Exécution", "Au moins un des 6 livrables obligatoires manque.", "Compléter la checklist de planification."],
            ["Validation DSI planification", "Validation DM absente.", "Le sponsor DM doit valider avant la DSI."],
            ["Clôture finale", "La chaîne demandeur -> DM -> DSI n'est pas complète.", "Valider chaque niveau sur les écrans dédiés."],
        ],
        [4.0, 5.5, 6.5],
    )

    add_heading(doc, "10. Workflow global synthétique", 1)
    add_number(doc, "Le demandeur crée la demande et joint le cahier des charges.")
    add_number(doc, "Le DM valide ou renvoie la demande.")
    add_number(doc, "La DSI valide la demande et crée le projet.")
    add_number(doc, "Le CP prépare l'analyse, la charte, la version signée et demande la validation charte.")
    add_number(doc, "Le DM valide la charte sur ValidationsProjet.")
    add_number(doc, "La DSI ou le RSIT délégué valident la charte sur ValidationsProjet.")
    add_number(doc, "Le CP valide la phase Analyse et passe en planification.")
    add_number(doc, "Le CP prépare la planification et charge les 6 livrables obligatoires.")
    add_number(doc, "Le DM valide la planification, puis la DSI ou le RSIT délégué passent le projet en exécution.")
    add_number(doc, "Le CP pilote exécution, UAT/MEP puis soumet la clôture.")
    add_number(doc, "Le demandeur, le DM puis la DSI valident la clôture finale.")

    add_heading(doc, "11. Conclusion opérationnelle", 1)
    add_body(
        doc,
        "Le workflow par rôle est désormais centré sur les écrans réellement actifs dans l'application. "
        "Les validations ne sont pas seulement documentaires : elles sont bloquées ou ouvertes par des "
        "contrôles serveur, des livrables obligatoires et des droits de rôle. Ce document doit être utilisé "
        "comme référence d'exploitation pour la DSI, les chefs de projet, les sponsors métiers et l'administration applicative."
    )
    return doc


def main() -> None:
    if TARGET.exists() and not BACKUP.exists():
        shutil.copy2(TARGET, BACKUP)

    doc = build_document()
    props = doc.core_properties
    props.title = "Workflow par rôle - Gestion Projets IT"
    props.subject = "Workflow applicatif actualisé"
    props.author = "OpenAI Codex"
    props.comments = "Document réécrit sur la base du workflow effectivement implémenté au 20/05/2026."
    props.created = datetime(2026, 5, 20, 10, 0, 0)
    props.modified = datetime(2026, 5, 20, 10, 0, 0)

    doc.save(TARGET)


if __name__ == "__main__":
    main()
