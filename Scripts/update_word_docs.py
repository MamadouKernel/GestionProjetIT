from __future__ import annotations

from collections import defaultdict
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
WORKFLOW_DOC = ROOT / "Workflow_Par_Role_GestionProjetsIT.docx"
DOCUMENTATION_DOC = ROOT / "Documentation_Complete_Application_GestionProjetsIT.docx"
VIEWS_ROOT = ROOT / "Views"
LOGO_PATH = ROOT / "wwwroot" / "images" / "LOGO_COTE_D_IVOIRE_TERMINAL.png"
TODAY = date(2026, 5, 21)


def save_with_fallback(doc: Document, path: Path) -> Path:
    try:
        doc.save(path)
        return path
    except PermissionError:
        fallback = path.with_name(path.stem + ".updated" + path.suffix)
        doc.save(fallback)
        return fallback


ROLE_MATRIX = [
    [
        "Demandeur",
        "Créer et suivre ses demandes, consulter ses projets, valider la clôture côté métier demandeur.",
        "Mes Demandes, Nouvelle Demande, Détails demande, Mes validations clôture, Tableau de bord, Aide.",
    ],
    [
        "Directeur Métier (Sponsor)",
        "Valider le besoin métier, valider la charte, valider la planification, valider la clôture DM.",
        "Validations DM, Validations Projet, Validations Clôture DM, Projets de ma direction, Historique & Traçabilité, Analytics direction.",
    ],
    [
        "Chef de Projet",
        "Préparer le dossier projet, renseigner les phases, produire les livrables natifs, piloter exécution/UAT/clôture.",
        "Mes projets, Détails projet, Charte projet, Charges, Planification, Exécution, UAT & MEP, Collaboration, Clôture.",
    ],
    [
        "DSI",
        "Valider transversalement le portefeuille, arbitrer la charte, valider la planification et la clôture finale.",
        "Validations DSI, Validations Projet, Portefeuille DSI, Liste des projets, Validations Clôture DSI, Dashboard, Analytics.",
    ],
    [
        "RSIT",
        "Consulter techniquement les projets, agir comme délégué DSI si une délégation active existe.",
        "Portefeuille, Liste des projets, Analyses, Validations Projet si délégation, Charges, Dashboard.",
    ],
    [
        "AdminIT",
        "Administrer l'application et conserver un accès total métier et technique.",
        "Administration complète, Autorisations / droits, Utilisateurs, Rôles, Directions, Services, Délégations, Paramètres, tous les écrans projet.",
    ],
]


ROLE_WORKFLOWS = [
    {
        "role": "Demandeur",
        "mission": "Le demandeur exprime le besoin, suit son instruction, corrige si nécessaire et intervient en clôture pour confirmer la fin du projet côté métier.",
        "screens": [
            (
                "Mes Demandes",
                "/DemandeProjet/Index",
                "L'écran liste les demandes créées par l'utilisateur, leur statut et les dernières actions. Il doit permettre la consultation rapide et la reprise des brouillons.",
                "Aucune saisie ici. Utiliser les filtres, ouvrir le détail et surveiller les notifications de retour DM/DSI.",
                "Affiche la chronologie, l'état de validation et les liens vers le détail ou la modification quand la demande est renvoyée.",
            ),
            (
                "Nouvelle Demande",
                "/DemandeProjet/Create",
                "Écran d'initialisation du besoin. Les contrôles empêchent l'envoi si les champs structurants sont absents.",
                "Renseigner le titre, le contexte, les objectifs, les parties prenantes, les bénéfices attendus, les impacts, l'urgence et les pièces justificatives.",
                "Création d'une demande en attente de validation DM et déclenchement des notifications.",
            ),
            (
                "Détail Demande",
                "/DemandeProjet/Details/{id}",
                "Affiche l'état courant de la demande et les décisions prises. Les actions disponibles changent selon le statut.",
                "En cas de retour, compléter les sections demandées puis réexpédier. Sinon écran de lecture.",
                "Permet de comprendre précisément ce qui bloque la demande et ce qui a déjà été validé.",
            ),
            (
                "Validations Clôture Demandeur",
                "/Projet/ListeValidationClotureDemandeur",
                "Liste les projets terminés qui attendent la confirmation de fin côté demandeur.",
                "Ouvrir chaque projet, vérifier le bilan et utiliser le bouton de validation ou de refus avec commentaire.",
                "Débloque la validation de clôture côté Directeur Métier.",
            ),
        ],
    },
    {
        "role": "Directeur Métier",
        "mission": "Le Directeur Métier est le sponsor. Il valide l'opportunité métier, la charte, la planification et la clôture métier.",
        "screens": [
            (
                "Validations DM",
                "/DemandeProjet/ListeValidationDM",
                "L'écran ne doit montrer que les demandes relevant du périmètre du DM. Les boutons ouvrent les modales de validation, correction ou refus.",
                "Lire le résumé, contrôler les pièces, demander correction si nécessaire, puis valider ou refuser avec commentaire motivé.",
                "La validation DM alimente le workflow initial et, si positive, ouvre la validation DSI.",
            ),
            (
                "Validations Projet",
                "/Projet/ValidationsProjet",
                "Liste les projets en phase Analyse dont la charte est prête à validation. Le bouton vert doit rester inactif tant que la charte signée complète n'est pas prête.",
                "Contrôler la charte, les signatures sponsor et chef de projet, consulter les pièces d'analyse, puis valider ou refuser la charte.",
                "Après validation DM, la validation DSI/RSIT devient possible.",
            ),
            (
                "Planification & validation",
                "/Projet/Details/{id}?tab=planification",
                "Le bloc Validation Directeur Métier présente un statut, les commentaires et les boutons de décision. Le DM ne saisit pas le planning, il le contrôle.",
                "Vérifier les livrables de planification, la cohérence du planning interactif, du budget et de la gouvernance. Cliquer Valider la planification ou Refuser avec justification.",
                "Conditionne l'ouverture du passage en exécution pour la DSI.",
            ),
            (
                "Validations Clôture DM",
                "/Projet/ListeValidationClotureDM",
                "Liste les clôtures déjà validées par le demandeur et encore en attente côté sponsor.",
                "Contrôler le bilan, les leçons apprises et le transfert au RUN avant validation.",
                "Débloque la validation finale DSI.",
            ),
        ],
    },
    {
        "role": "Chef de Projet",
        "mission": "Le Chef de Projet pilote de bout en bout : analyse, charte, planning natif, exécution, UAT/MEP, collaboration, charges et clôture.",
        "screens": [
            (
                "Détail projet - Analyse",
                "/Projet/Details/{id}?tab=analyse",
                "L'écran doit exposer le contexte, les objectifs, les documents d'analyse, la situation de la charte et les alertes de blocage.",
                "Compléter les informations d'analyse, déposer le cahier d'analyse technique et la note de cadrage, renseigner les risques initiaux, puis générer la charte PDF.",
                "La phase ne peut être validée que si Charte projet, Charte projet signée, validation DM et validation DSI sont réunies.",
            ),
            (
                "Charte Projet",
                "/Projet/CharteProjet/{id}",
                "Écran de rédaction de la charte, de dépôt de la version signée et de gestion des signatures sponsor / chef de projet.",
                "Renseigner les sections de charte, sauvegarder, générer la version PDF, déposer la version signée, cocher Sponsor / Directeur Métier et Chef de Projet, puis enregistrer les signatures.",
                "Alimente les validations projet et remet à zéro les validations DM/DSI si la charte signée change.",
            ),
            (
                "Planification",
                "/Projet/Details/{id}?tab=planification",
                "Le CP prépare le planning natif. Les blocs RACI, communication, budget et PV de kick-off sont désormais saisis directement dans l'application.",
                "Ajouter les tâches dans le Gantt, renseigner jalons, WBS, RACI, communication, budget et PV de kick-off, enregistrer puis générer les livrables Excel officiels.",
                "Les livrables Planning détaillé, WBS, Matrice RACI, Schéma de communication, Budget prévisionnel et PV de kick-off sont produits avec le logo.",
            ),
            (
                "Exécution",
                "/Projet/Details/{id}?tab=execution",
                "Écran de suivi opérationnel. Il doit permettre d'actualiser l'avancement réel, les blocages, les décisions et les écarts planning/budget.",
                "Renseigner l'avancement, les dates réelles, les actions réalisées, les actions à venir, les blocages, la justification des retards, le besoin de changement et la synthèse de charges.",
                "Fournit la base de pilotage hebdomadaire et les indicateurs remontés au portefeuille.",
            ),
            (
                "Charges",
                "/Projet/Charges/{id}",
                "Cockpit de saisie des charges réelles et prévues par semaine. L'écran pilote la capacité et les alertes de surcharge.",
                "Saisir semaine par semaine les heures prévues, les heures réelles et les commentaires par ressource / activité.",
                "Alimente les indicateurs charges/capacité et le suivi hebdomadaire.",
            ),
            (
                "UAT & MEP",
                "/Projet/Details/{id}?tab=uat",
                "Bloc de recette, anomalies, mise en production, rollback et hypercare. Les validations s'appuient sur les données renseignées ici.",
                "Renseigner les campagnes de test, cas de test, anomalies, plan de MEP, prérequis, rollback, résultat de MEP, incidents et période d'hypercare.",
                "Conditionne la clôture technique du projet.",
            ),
            (
                "Clôture",
                "/Projet/Details/{id}?tab=cloture",
                "Écran de bilan final. Il doit résumer résultats, statut final, transfert au RUN et validations encore attendues.",
                "Renseigner le statut final, les commentaires, le transfert RUN, les accès, la documentation, l'état du support et les leçons apprises.",
                "Déclenche les validations Demandeur, DM puis DSI.",
            ),
        ],
    },
    {
        "role": "DSI / RSIT délégué",
        "mission": "La DSI valide transversalement les demandes, la charte, la planification et la clôture. Le RSIT n'agit en lieu et place de la DSI que si une délégation est active.",
        "screens": [
            (
                "Validations DSI",
                "/DemandeProjet/ListeValidationDSI",
                "L'écran ne doit montrer que les demandes déjà validées par le DM. Les actions sont tracées dans l'historique DSI.",
                "Contrôler le besoin, les impacts, l'alignement portefeuille, puis valider, refuser ou renvoyer.",
                "Crée automatiquement le projet quand la demande est approuvée.",
            ),
            (
                "Validations Projet",
                "/Projet/ValidationsProjet",
                "L'écran présente les mêmes projets que le DM, mais la décision DSI n'est possible qu'après validation DM.",
                "Vérifier la charte, les signatures, les livrables d'analyse et les remarques métier. Cliquer Valider ou Refuser.",
                "Autorise ensuite la validation finale de la phase Analyse par le CP.",
            ),
            (
                "Planification",
                "/Projet/Details/{id}?tab=planification",
                "Le bloc Validation DSI doit rester sur En attente DM tant que le sponsor n'a pas validé. Après validation DM, la DSI peut faire passer en exécution.",
                "Vérifier le planning, la gouvernance, les six livrables obligatoires et le Go/No-Go. Cliquer Passer en exécution.",
                "Bascule la phase en Exécution et ouvre le pilotage opérationnel.",
            ),
            (
                "Validations Clôture DSI",
                "/Projet/ListeValidationClotureDSI",
                "Liste finale des projets prêts à clôture DSI après validations demandeur et sponsor.",
                "Contrôler le bilan, les incidents résiduels, le transfert RUN et valider la clôture finale.",
                "Achève le cycle projet.",
            ),
        ],
    },
    {
        "role": "AdminIT",
        "mission": "AdminIT garde un accès total. Il paramètre l'application, les utilisateurs et la matrice des droits sans être limité par les rôles métier.",
        "screens": [
            (
                "Utilisateurs",
                "/Admin/Users",
                "Gestion des comptes, affectations de rôles et réinitialisations. Les formulaires doivent appliquer la politique mot de passe.",
                "Créer ou modifier un utilisateur, attribuer direction, service et rôles, réinitialiser le mot de passe si nécessaire.",
                "Provisionne les comptes et régit l'accès global.",
            ),
            (
                "Délégations",
                "/Admin/Delegations",
                "Administration des délégations DSI et Chef de Projet. Les périodes d'activité pilotent les remplacements réels dans les workflows.",
                "Créer ou modifier une délégation en renseignant le titulaire, le remplaçant, le périmètre et la période d'effet.",
                "Active les droits de validation délégués en temps réel.",
            ),
            (
                "Autorisations / Droits",
                "/Autorisations/Index",
                "Matrice de permissions réellement branchée sur le menu et sur le blocage backend.",
                "Activer ou désactiver une route par rôle, puis réinitialiser si besoin pour revenir au modèle par défaut.",
                "Pilote l'affichage des menus et l'accès réel par URL.",
            ),
        ],
    },
]


PHASE_GUIDES = [
    (
        "Demande initiale",
        "Le cycle démarre par la saisie d'une demande structurée par le demandeur. Le Directeur Métier valide en premier, puis la DSI.",
        [
            "Écrans : /DemandeProjet/Create, /DemandeProjet/Details/{id}, /DemandeProjet/ListeValidationDM, /DemandeProjet/ListeValidationDSI.",
            "Livrables : pièces jointes justificatives si besoin.",
            "Blocages : champs obligatoires absents, retour pour correction, refus DM ou DSI.",
            "Sortie de phase : création automatique du projet après validation DSI.",
        ],
    ),
    (
        "Analyse & Clarification",
        "La phase Analyse sert à cadrer le besoin, produire la charte, collecter les documents d'analyse et faire valider la charte par DM puis DSI/RSIT.",
        [
            "Écrans : /Projet/Details/{id}?tab=analyse, /Projet/CharteProjet/{id}, /Projet/ValidationsProjet.",
            "Livrables : Cahier d'analyse technique, Note de cadrage, Charte projet (générée), Charte projet signée.",
            "Blocages : charte non générée, charte signée absente ou incomplète, validation DM absente, validation DSI absente.",
            "Sortie de phase : bouton Valider la phase Analyse activé uniquement quand tous les prérequis sont satisfaits.",
        ],
    ),
    (
        "Planification & Validation",
        "Le planning détaillé devient natif. Le CP saisit tâches, RACI, communication, budget et PV de kick-off dans l'application, puis génère les livrables officiels avec logo.",
        [
            "Écrans : /Projet/Details/{id}?tab=planification.",
            "Livrables générés : Planning détaillé, WBS, Matrice RACI, Schéma de communication, Budget prévisionnel, PV de kick-off (Excel brandés).",
            "Blocages : aucune tâche enregistrée, livrables obligatoires absents, validation DM non obtenue, validation DSI non obtenue.",
            "Sortie de phase : passage en Exécution par la DSI ou le RSIT délégué après validation sponsor.",
        ],
    ),
    (
        "Exécution",
        "La phase Exécution couvre le pilotage opérationnel, l'avancement, les charges, les changements, les risques et les livrables d'avancement.",
        [
            "Écrans : /Projet/Details/{id}?tab=execution, /Projet/Charges/{id}.",
            "Livrables : rapports d'avancement, annexes, charges hebdomadaires, décisions et justifications.",
            "Blocages : données d'avancement incomplètes, charges manquantes, anomalies critiques non traitées.",
            "Sortie de phase : préparation de la recette et de la MEP.",
        ],
    ),
    (
        "UAT & MEP",
        "La phase UAT & MEP organise les tests, les anomalies, la mise en production, le rollback et l'hypercare.",
        [
            "Écrans : /Projet/Details/{id}?tab=uat, partiel Cas de tests.",
            "Livrables : campagnes de test, cas de test, plan de MEP, rollback, compte-rendu de recette, éléments hypercare.",
            "Blocages : anomalies critiques ouvertes, MEP non préparée, hypercare non terminée.",
            "Sortie de phase : ouverture de la clôture.",
        ],
    ),
    (
        "Clôture",
        "La clôture consolide le bilan et orchestre les validations demandeur, sponsor puis DSI.",
        [
            "Écrans : /Projet/Details/{id}?tab=cloture, /Projet/ListeValidationClotureDemandeur, /Projet/ListeValidationClotureDM, /Projet/ListeValidationClotureDSI.",
            "Livrables : bilan de clôture, documentation RUN, preuves de transfert, commentaires finaux.",
            "Blocages : validations manquantes, transfert RUN incomplet, statut final non renseigné.",
            "Sortie de phase : projet définitivement clôturé.",
        ],
    ),
]


DETAILED_SCREENS = [
    {
        "name": "Connexion",
        "route": "/Account/Login",
        "roles": "Tous les utilisateurs",
        "purpose": "Entrer dans l'application via compte local autorisé ou via Azure AD selon la configuration.",
        "behavior": "L'écran doit afficher clairement le mode d'authentification disponible, bloquer les tentatives excessives et rediriger vers le tableau de bord adapté au rôle.",
        "fill": "Saisir identifiant et mot de passe si l'authentification locale est autorisée. Sinon utiliser le bouton SSO.",
        "outputs": "Création de session, chargement des menus autorisés, journalisation de connexion.",
        "blockers": "Compte inactif, mot de passe invalide, authentification locale désactivée, trop de tentatives.",
    },
    {
        "name": "Tableau de bord",
        "route": "/Home/Index",
        "roles": "Tous, avec contenu différent selon le rôle",
        "purpose": "Fournir un cockpit d'entrée : KPI, validations en attente, alertes, échéances et synthèse portefeuille.",
        "behavior": "L'écran doit privilégier l'information utile au rôle connecté. Les cartes ouvrent les écrans de détail appropriés.",
        "fill": "Pas de saisie. Utiliser les liens d'action pour aller vers les validations ou les projets critiques.",
        "outputs": "Vue consolidée des priorités immédiates.",
        "blockers": "Aucun, hors absence de données.",
    },
    {
        "name": "Centre de notifications",
        "route": "/Notification/Index",
        "roles": "Tous",
        "purpose": "Centraliser les événements métier et techniques à lire.",
        "behavior": "Chaque notification doit pouvoir être marquée comme lue sans sortir de l'écran. Le bouton Tout marquer comme lu doit rester dans le flux de la page.",
        "fill": "Pas de saisie. Filtrer, ouvrir la notification liée puis marquer comme lue.",
        "outputs": "Réduction du compteur de notifications, traçabilité de lecture.",
        "blockers": "Aucun.",
    },
    {
        "name": "Mes Demandes",
        "route": "/DemandeProjet/Index",
        "roles": "Demandeur, AdminIT si accès complet",
        "purpose": "Lister les demandes créées par l'utilisateur et leur statut.",
        "behavior": "Les demandes doivent être filtrables, consultables et éventuellement modifiables si elles sont renvoyées.",
        "fill": "Pas de saisie. Utiliser les filtres, ouvrir la demande ou créer une nouvelle demande.",
        "outputs": "Suivi lisible du portefeuille personnel de demandes.",
        "blockers": "Aucun.",
    },
    {
        "name": "Nouvelle Demande",
        "route": "/DemandeProjet/Create",
        "roles": "Demandeur, AdminIT si besoin métier",
        "purpose": "Créer un besoin projet complet et directement exploitable par le DM puis la DSI.",
        "behavior": "Le formulaire doit guider pas à pas, prévenir les doublons et empêcher l'envoi si le minimum métier n'est pas renseigné.",
        "fill": "Titre, contexte, objectifs, bénéfices, parties prenantes, direction, impacts, urgence, pièces jointes et toute donnée de cadrage initiale.",
        "outputs": "Demande enregistrée, horodatée et soumise au DM.",
        "blockers": "Champs obligatoires vides, doublon métier confirmé sans justification, pièces jointes invalides.",
    },
    {
        "name": "Validation DM de la demande",
        "route": "/DemandeProjet/ListeValidationDM",
        "roles": "Directeur Métier, AdminIT",
        "purpose": "Décider du go / no-go métier sur les demandes initiales.",
        "behavior": "Le DM doit voir uniquement son périmètre et disposer de trois issues : valider, corriger, rejeter.",
        "fill": "Renseigner un commentaire explicatif pour tout refus ou demande de correction.",
        "outputs": "Demande transmise à la DSI, renvoyée au demandeur ou rejetée.",
        "blockers": "Aucun technique ; blocage purement décisionnel.",
    },
    {
        "name": "Validation DSI de la demande",
        "route": "/DemandeProjet/ListeValidationDSI",
        "roles": "DSI, RSIT délégué, AdminIT",
        "purpose": "Arbitrer les demandes déjà validées par le sponsor.",
        "behavior": "La DSI ne doit voir que les demandes validées DM. Les décisions doivent être tracées et créer le projet en cas d'approbation.",
        "fill": "Commentaire de décision, éventuelles consignes de cadrage.",
        "outputs": "Projet créé automatiquement ou demande refusée/renvoyée.",
        "blockers": "Validation DM absente.",
    },
    {
        "name": "Détail projet - Synthèse",
        "route": "/Projet/Details/{id}?tab=synthese",
        "roles": "Acteurs du projet autorisés",
        "purpose": "Concentrer la vision de haut niveau du projet : phase, RAG, équipe, risques, anomalies, jalons et validations.",
        "behavior": "Les cartes doivent être lues comme un cockpit, avec liens vers les écrans d'action (analyse, risques, anomalies, collaboration).",
        "fill": "Peu de saisie directe ; c'est principalement une vue de synthèse alimentée par les autres onglets.",
        "outputs": "Vue de pilotage rapide.",
        "blockers": "Aucun.",
    },
    {
        "name": "Détail projet - Analyse",
        "route": "/Projet/Details/{id}?tab=analyse",
        "roles": "CP, DSI, RSIT, AdminIT, lecture contrôlée pour DM",
        "purpose": "Cadrer le besoin, déposer les pièces d'analyse, suivre la charte et préparer le passage en planification.",
        "behavior": "L'écran doit lister clairement ce qui manque : charte projet, charte signée, validation DM, validation DSI. Les boutons de validation finale restent bloqués tant que les prérequis ne sont pas satisfaits.",
        "fill": "Contexte, objectifs, documents d'analyse obligatoires, autres livrables d'analyse, risques et membres si nécessaire.",
        "outputs": "Dossier d'analyse complet et message de blocage explicite en cas d'élément manquant.",
        "blockers": "Charte projet non générée, charte signée absente, signatures manquantes, validations DM/DSI absentes.",
    },
    {
        "name": "Charte Projet",
        "route": "/Projet/CharteProjet/{id}",
        "roles": "CP, DSI, RSIT, AdminIT, consultation DM",
        "purpose": "Rédiger la charte, générer la charte projet officielle, déposer la version signée et gérer les signatures.",
        "behavior": "L'écran doit distinguer rédaction de charte, dossier de signature et version signée déposée. Les cases Sponsor / Chef de Projet doivent pouvoir être sauvegardées séparément.",
        "fill": "Sections de charte, fichier signé si disponible, cases Sponsor / Directeur Métier et Chef de Projet.",
        "outputs": "Charte projet générée, charte signée liée au projet, signatures persistées et validations réinitialisées si changement.",
        "blockers": "Phase hors Analyse pour certaines actions, fichier absent si dépôt signé requis.",
    },
    {
        "name": "Validation de la charte projet",
        "route": "/Projet/ValidationsProjet",
        "roles": "DM, DSI, RSIT délégué, AdminIT",
        "purpose": "Valider la charte une fois le dossier de signature complet.",
        "behavior": "Le bouton vert doit être clairement bloqué tant que la charte signée complète n'est pas présente. Le bouton œil ouvre le détail en lecture.",
        "fill": "Commentaire de refus ou d'observation le cas échéant.",
        "outputs": "Validation DM puis validation DSI. Les statuts doivent changer immédiatement.",
        "blockers": "Charte signée incomplète, validation DM manquante pour la DSI, délégation DSI absente pour le RSIT.",
    },
    {
        "name": "Détail projet - Planification",
        "route": "/Projet/Details/{id}?tab=planification",
        "roles": "CP, DSI, AdminIT, validation DM/DSI en lecture décisionnelle",
        "purpose": "Construire le planning de référence et produire les livrables de planification sans quitter l'application.",
        "behavior": "Le planning interactif doit permettre l'ajout de tâches, la sauvegarde, puis la génération des livrables officiels avec logo. Les validations sponsor et DSI restent visibles en tête de l'écran.",
        "fill": "Tâches Gantt, jalons, responsables, dépendances, RACI, communication, budget, PV de kick-off et commentaires de validation.",
        "outputs": "Livrables Planning détaillé, WBS, Matrice RACI, Schéma de communication, Budget prévisionnel et PV de kick-off en Excel brandé.",
        "blockers": "Aucune tâche enregistrée, dossier incomplet, validation sponsor absente, validation DSI absente.",
    },
    {
        "name": "Détail projet - Exécution",
        "route": "/Projet/Details/{id}?tab=execution",
        "roles": "CP, DSI, RSIT, AdminIT",
        "purpose": "Piloter le projet en cours d'exécution.",
        "behavior": "Les données saisies ici doivent alimenter immédiatement la synthèse, les alertes et le dashboard.",
        "fill": "Avancement, dates réelles, actions réalisées et à venir, blocages, décisions, justification des retards, changement requis, synthèse des charges.",
        "outputs": "Vision opérationnelle et matière pour les points hebdomadaires.",
        "blockers": "Peu de blocages techniques ; qualité des données exigée pour pilotage fiable.",
    },
    {
        "name": "Charges / Capacité",
        "route": "/Projet/Charges/{id}",
        "roles": "CP, DSI, RSIT, AdminIT",
        "purpose": "Suivre les charges prévues et réelles du projet.",
        "behavior": "La grille hebdomadaire doit permettre une saisie rapide par semaine et remonter les alertes de surcharge.",
        "fill": "Heures prévues, heures réelles, commentaires et activités par ressource / semaine.",
        "outputs": "Cockpit charges/capacité, alertes, indicateurs portefeuille.",
        "blockers": "Aucun, sauf droits insuffisants.",
    },
    {
        "name": "Détail projet - UAT & MEP",
        "route": "/Projet/Details/{id}?tab=uat",
        "roles": "CP, DSI, RSIT, AdminIT",
        "purpose": "Gérer tests, anomalies, mise en production, rollback et hypercare.",
        "behavior": "L'écran doit distinguer les prérequis, la recette, la mise en production, l'hypercare et l'état des anomalies.",
        "fill": "Campagnes de test, cas de test, anomalies, utilisateurs testeurs, plan de MEP, rollback, incidents, période d'hypercare.",
        "outputs": "Décision de mise en production et traçabilité de recette.",
        "blockers": "Anomalies critiques, éléments de MEP manquants, hypercare non finalisée.",
    },
    {
        "name": "Détail projet - Collaboration",
        "route": "/Projet/Details/{id}?tab=collaboration",
        "roles": "CP, DSI, RSIT, AdminIT",
        "purpose": "Structurer l'équipe projet, les interactions et le partage d'information.",
        "behavior": "Les membres, rôles et échanges doivent être visibles et éditables sans sortir du projet.",
        "fill": "Participants, responsabilités, observations collaboratives et éléments de gouvernance.",
        "outputs": "Vision claire de l'équipe et de la coordination.",
        "blockers": "Aucun.",
    },
    {
        "name": "Détail projet - Clôture",
        "route": "/Projet/Details/{id}?tab=cloture",
        "roles": "CP, DSI, AdminIT, lecture demandeur/DM selon droits",
        "purpose": "Préparer le bilan final et le transfert au RUN.",
        "behavior": "L'écran doit afficher l'état des validations demandeur, sponsor et DSI, ainsi que les éléments de transfert manquants.",
        "fill": "Statut final, commentaire final, transfert RUN, accès, documentation, support, incidents résiduels, leçons apprises.",
        "outputs": "Bilan de clôture exploitable et lancement du workflow de validations finales.",
        "blockers": "Éléments RUN non renseignés, validations intermédiaires absentes.",
    },
    {
        "name": "Portefeuille Projets",
        "route": "/Projet/Portefeuille",
        "roles": "DSI, RSIT, DM selon filtrage, AdminIT",
        "purpose": "Piloter le portefeuille consolidé des projets.",
        "behavior": "L'écran doit donner une vue filtrable des projets, de leur santé, des exports et des validations en attente.",
        "fill": "Pas de saisie. Utiliser filtres, exports PDF/Excel et liens vers les projets.",
        "outputs": "Pilotage portefeuille, rapport DSI/DG, portefeuille PDF/Excel.",
        "blockers": "Aucun.",
    },
    {
        "name": "Historique & Traçabilité",
        "route": "/Projet/HistoriqueDM",
        "roles": "DM, DSI, AdminIT, autres selon droits",
        "purpose": "Visualiser l'historique complet d'un portefeuille de projets ou d'une direction.",
        "behavior": "Les fiches projet sont pliables pour alléger la lecture. Chaque bloc doit détailler cycle de vie, validations, risques et anomalies.",
        "fill": "Pas de saisie. Ouvrir / replier les projets et naviguer vers les détails.",
        "outputs": "Traçabilité consolidée.",
        "blockers": "Aucun.",
    },
    {
        "name": "Administration des utilisateurs",
        "route": "/Admin/Users",
        "roles": "AdminIT",
        "purpose": "Créer, modifier, désactiver les utilisateurs et gérer leurs rôles.",
        "behavior": "Les formulaires doivent appliquer la politique mot de passe et les règles d'affectation rôle/direction/service.",
        "fill": "Identité, email, login, rôle(s), direction, service, statut, mot de passe si création ou reset.",
        "outputs": "Comptes opérationnels, rôles mis à jour, messages de succès/erreur explicites.",
        "blockers": "Mot de passe non conforme, doublons, rôle invalide.",
    },
    {
        "name": "Matrice des autorisations",
        "route": "/Autorisations/Index",
        "roles": "AdminIT",
        "purpose": "Piloter les droits par rôle sur les routes applicatives.",
        "behavior": "Toute modification doit avoir un effet réel sur le menu et l'accès backend. Le bouton Réinitialiser doit recharger le standard par défaut.",
        "fill": "Activer / désactiver les cases par rôle et par vue.",
        "outputs": "Permissions persistées, menus ajustés, accès serveur bloqués ou ouverts.",
        "blockers": "Aucun pour AdminIT.",
    },
]


VIEW_DESCRIPTIONS = {
    "Account/AccessDenied.cshtml": ("Écran", "Page de refus d'accès lorsque l'utilisateur n'a pas la permission demandée.", "Aucune saisie ; retour à une page autorisée."),
    "Account/DemandeAcces.cshtml": ("Écran", "Demande d'accès complémentaire ou de création de compte local.", "Renseigner les informations de demande d'accès et le besoin métier."),
    "Account/Inscription.cshtml": ("Écran", "Onboarding / inscription lorsque le workflow local de création de compte est autorisé.", "Renseigner identité, coordonnées, rattachement et justification."),
    "Account/Login.cshtml": ("Écran", "Point d'entrée authentification locale / SSO.", "Saisir identifiant et mot de passe si mode local."),
    "Account/Profil.cshtml": ("Écran", "Profil utilisateur, sécurité et activité.", "Mettre à jour les informations personnelles ou le mot de passe."),
    "Admin/Delegations.cshtml": ("Écran", "Gestion des délégations DSI.", "Créer ou modifier une délégation active."),
    "Admin/DelegationsChefProjet.cshtml": ("Écran", "Gestion des délégations Chef de Projet.", "Définir CP titulaire, remplaçant, période et périmètre."),
    "Admin/DemandesCreationCompte.cshtml": ("Écran", "Traitement des demandes de création de compte.", "Valider, refuser ou corriger une demande de compte."),
    "Admin/Directions.cshtml": ("Écran", "Référentiel des directions.", "Créer, modifier ou désactiver une direction."),
    "Admin/GererRoles.cshtml": ("Écran", "Affectation détaillée des rôles utilisateur.", "Cocher ou décocher les rôles autorisés."),
    "Admin/ImportUsers.cshtml": ("Écran", "Import massif d'utilisateurs.", "Déposer un fichier d'import et contrôler les résultats."),
    "Admin/ListeRoles.cshtml": ("Écran", "Catalogue des rôles applicatifs.", "Consulter ou ajuster le périmètre des rôles."),
    "Admin/Parametres.cshtml": ("Écran", "Paramètres globaux applicatifs.", "Modifier les paramètres d'exploitation autorisés."),
    "Admin/Services.cshtml": ("Écran", "Référentiel des services.", "Créer, modifier ou désactiver un service."),
    "Admin/Users.cshtml": ("Écran", "Administration des utilisateurs.", "Créer, modifier, réinitialiser, attribuer les rôles."),
    "Admin/_ModalDelegationChefProjet.cshtml": ("Modal", "Fenêtre de saisie des délégations CP.", "Renseigner délégant, délégataire, dates et commentaire."),
    "Admin/_ModalDelegationDSI.cshtml": ("Modal", "Fenêtre de saisie des délégations DSI.", "Renseigner les paramètres de délégation."),
    "Aide/Index.cshtml": ("Écran", "Centre d'aide rôle par rôle.", "Choisir le guide correspondant au rôle."),
    "Aide/_GuideAdminIT.cshtml": ("Guide", "Guide AdminIT.", "Lecture seule."),
    "Aide/_GuideChefProjet.cshtml": ("Guide", "Guide Chef de Projet.", "Lecture seule."),
    "Aide/_GuideDemandeur.cshtml": ("Guide", "Guide Demandeur.", "Lecture seule."),
    "Aide/_GuideDirecteurMetier.cshtml": ("Guide", "Guide Directeur Métier.", "Lecture seule."),
    "Aide/_GuideDSI.cshtml": ("Guide", "Guide DSI.", "Lecture seule."),
    "Aide/_GuideResponsableSolutionsIT.cshtml": ("Guide", "Guide RSIT.", "Lecture seule."),
    "Autorisations/Index.cshtml": ("Écran", "Matrice des droits réellement branchée sur l'application.", "Activer ou désactiver les permissions par rôle."),
    "AzureAuth/DemanderAcces.cshtml": ("Écran", "Demande d'accès dans le flux Azure AD.", "Renseigner identité et motif d'accès."),
    "Dashboard/Index.cshtml": ("Écran", "Dashboard secondaire / analytique.", "Lecture seule avec filtres."),
    "DemandeProjet/Create.cshtml": ("Écran", "Création d'une nouvelle demande projet.", "Compléter le formulaire métier et les pièces."),
    "DemandeProjet/Details.cshtml": ("Écran", "Détail d'une demande projet.", "Lecture ou correction selon le statut."),
    "DemandeProjet/Edit.cshtml": ("Écran", "Modification d'une demande renvoyée.", "Corriger les informations demandées puis renvoyer."),
    "DemandeProjet/HistoriqueActionsDM.cshtml": ("Écran", "Historique des actions DM sur les demandes.", "Lecture seule avec filtres."),
    "DemandeProjet/HistoriqueValidationsDSI.cshtml": ("Écran", "Historique des validations DSI.", "Lecture seule avec filtres."),
    "DemandeProjet/Index.cshtml": ("Écran", "Mes demandes du demandeur.", "Lecture, filtrage et accès aux détails."),
    "DemandeProjet/ListeValidationDM.cshtml": ("Écran", "File d'attente de validation sponsor.", "Valider, corriger ou refuser."),
    "DemandeProjet/ListeValidationDSI.cshtml": ("Écran", "File d'attente de validation DSI.", "Valider, renvoyer ou refuser."),
    "DemandeProjet/VerificationDoublons.cshtml": ("Écran", "Aide à la détection des doublons de demandes.", "Comparer et confirmer la poursuite."),
    "DemandeProjet/_ModalCorrectionDM.cshtml": ("Modal", "Motif de correction DM.", "Saisir la demande de correction."),
    "DemandeProjet/_ModalRejeterDM.cshtml": ("Modal", "Refus DM.", "Saisir le motif de refus."),
    "DemandeProjet/_ModalRejeterDSI.cshtml": ("Modal", "Refus DSI.", "Saisir le motif de refus."),
    "DemandeProjet/_ModalRenvoyerDemandeur.cshtml": ("Modal", "Renvoi au demandeur.", "Saisir les corrections attendues."),
    "DemandeProjet/_ModalRenvoyerDM.cshtml": ("Modal", "Renvoi au sponsor.", "Saisir la raison du renvoi."),
    "DemandeProjet/_ModalValiderDM.cshtml": ("Modal", "Confirmation de validation DM.", "Confirmer et éventuellement commenter."),
    "DemandeProjet/_ModalValiderDSI.cshtml": ("Modal", "Confirmation de validation DSI.", "Confirmer et éventuellement commenter."),
    "DemandesAcces/Index.cshtml": ("Écran", "Traitement des demandes d'accès applicatives.", "Valider ou rejeter les demandes."),
    "Document/Preview.cshtml": ("Écran", "Prévisualisation sécurisée des documents.", "Lecture seule, téléchargement éventuel."),
    "Home/Index.cshtml": ("Écran", "Cockpit principal de l'application.", "Lecture seule et navigation."),
    "Home/Privacy.cshtml": ("Écran", "Politique / informations de confidentialité.", "Lecture seule."),
    "Notification/Index.cshtml": ("Écran", "Centre de notifications.", "Marquer les notifications comme lues."),
    "Projet/Charges.cshtml": ("Écran", "Saisie et suivi des charges du projet.", "Renseigner les heures prévues et réelles."),
    "Projet/CharteProjet.cshtml": ("Écran", "Rédaction et signature de la charte.", "Compléter la charte, déposer la version signée et les signatures."),
    "Projet/Details.cshtml": ("Écran", "Conteneur central à onglets du projet.", "Naviguer entre les onglets et saisir les données de phase."),
    "Projet/FicheProjet.cshtml": ("Écran", "Fiche projet éditoriale / documentaire.", "Compléter ou consulter la fiche de pilotage."),
    "Projet/HistoriqueDM.cshtml": ("Écran", "Historique portefeuille / direction avec fiches repliables.", "Lecture seule, ouverture / fermeture des projets."),
    "Projet/Index.cshtml": ("Écran", "Liste des projets.", "Lecture, filtrage et accès au détail."),
    "Projet/ListeValidationClotureDemandeur.cshtml": ("Écran", "Clôtures en attente côté demandeur.", "Valider ou refuser la clôture."),
    "Projet/ListeValidationClotureDM.cshtml": ("Écran", "Clôtures en attente côté sponsor.", "Valider ou refuser la clôture."),
    "Projet/ListeValidationClotureDSI.cshtml": ("Écran", "Clôtures en attente côté DSI.", "Valider ou refuser la clôture finale."),
    "Projet/Portefeuille.cshtml": ("Écran", "Portefeuille projets et exports.", "Lecture, filtres et exports."),
    "Projet/SignatureCharte.cshtml": ("Écran", "Vue de signature / suivi de charte.", "Lecture ou action selon le workflow."),
    "Projet/ValidationsProjet.cshtml": ("Écran", "Validation de la charte projet.", "Valider ou refuser la charte."),
    "Projet/_AjouterAnomalieModal.cshtml": ("Modal", "Ajout d'une anomalie.", "Renseigner gravité, description, responsable."),
    "Projet/_AjouterMembreModal.cshtml": ("Modal", "Ajout d'un membre projet.", "Renseigner personne, rôle et commentaire."),
    "Projet/_AjouterRisqueModal.cshtml": ("Modal", "Ajout d'un risque.", "Renseigner libellé, impact, probabilité et mitigation."),
    "Projet/_DossierSignature.cshtml": ("Partiel", "Suivi du dossier de signature.", "Consulter / déclencher les actions de signature."),
    "Projet/_ModifierMembreModal.cshtml": ("Modal", "Modification d'un membre projet.", "Ajuster rôle ou informations du membre."),
    "Projet/_ModifierRisqueModal.cshtml": ("Modal", "Modification d'un risque.", "Ajuster le risque et son plan de mitigation."),
    "Projet/_ProjetAnalyse.cshtml": ("Onglet", "Onglet Analyse du détail projet.", "Renseigner cadrage, documents d'analyse et état charte."),
    "Projet/_ProjetCasTests.cshtml": ("Partiel", "Gestion des cas de test UAT.", "Créer, exécuter ou mettre à jour les cas de test."),
    "Projet/_ProjetCloture.cshtml": ("Onglet", "Onglet Clôture du détail projet.", "Renseigner bilan final et transfert RUN."),
    "Projet/_ProjetCollaboration.cshtml": ("Onglet", "Onglet Collaboration du détail projet.", "Renseigner équipe, coordination et éléments collaboratifs."),
    "Projet/_ProjetExecution.cshtml": ("Onglet", "Onglet Exécution du détail projet.", "Renseigner avancement, blocages, décisions et dates réelles."),
    "Projet/_ProjetHistorique.cshtml": ("Onglet", "Onglet Historique du détail projet.", "Consulter la traçabilité."),
    "Projet/_ProjetPlanification.cshtml": ("Onglet", "Onglet Planification du détail projet.", "Renseigner planning natif, RACI, communication, budget, PV."),
    "Projet/_ProjetSynthese.cshtml": ("Onglet", "Onglet Synthèse du détail projet.", "Lecture consolidée du projet."),
    "Projet/_ProjetUAT.cshtml": ("Onglet", "Onglet UAT & MEP du détail projet.", "Renseigner recette, anomalies, MEP et hypercare."),
    "Projet/_UploadLivrableModal.cshtml": ("Modal", "Upload / remplacement d'un livrable projet.", "Choisir type de livrable, fichier et commentaire."),
    "Projet/_ValidationCharte.cshtml": ("Partiel", "Bloc de validation de charte.", "Lecture / validation selon rôle."),
    "Shared/Error.cshtml": ("Écran", "Page d'erreur standard.", "Lecture seule."),
    "Shared/_Layout.cshtml": ("Shell", "Structure commune : topbar, sidebar, notifications, footer.", "Aucune saisie directe."),
    "Shared/_Pagination.cshtml": ("Partiel", "Composant de pagination.", "Utiliser pour naviguer entre pages."),
    "Shared/_ValidationScriptsPartial.cshtml": ("Technique", "Scripts de validation client.", "Aucune action utilisateur."),
    "Shared/_ValidationSummary.cshtml": ("Partiel", "Résumé des erreurs de validation.", "Lecture des erreurs avant correction."),
    "Shared/Components/SidebarMenu/Default.cshtml": ("Composant", "Menu latéral piloté par la matrice de permissions.", "Aucune saisie ; navigation."),
}


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def apply_doc_defaults(doc: Document, title: str) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.6)
    section.left_margin = Cm(1.9)
    section.right_margin = Cm(1.7)

    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10.5)
    styles["Heading 1"].font.name = "Aptos Display"
    styles["Heading 1"].font.size = Pt(18)
    styles["Heading 1"].font.color.rgb = RGBColor(31, 45, 61)
    styles["Heading 2"].font.name = "Aptos Display"
    styles["Heading 2"].font.size = Pt(14)
    styles["Heading 2"].font.color.rgb = RGBColor(49, 71, 110)
    styles["Heading 3"].font.name = "Aptos Display"
    styles["Heading 3"].font.size = Pt(11.5)
    styles["Heading 3"].font.color.rgb = RGBColor(59, 76, 109)

    doc.core_properties.title = title
    doc.core_properties.author = "OpenAI Codex"
    doc.core_properties.subject = "Gestion Projets IT"
    doc.core_properties.comments = "Document régénéré depuis le code réel de l'application."


def add_header(section, title: str) -> None:
    header = section.header
    header.is_linked_to_previous = False
    paragraph = header.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if LOGO_PATH.exists():
        run = paragraph.add_run()
        run.add_picture(str(LOGO_PATH), width=Inches(0.38))
        paragraph.add_run("  ")
    title_run = paragraph.add_run(title)
    title_run.bold = True
    title_run.font.size = Pt(9)
    title_run.font.color.rgb = RGBColor(31, 45, 61)

    footer = section.footer
    footer.is_linked_to_previous = False
    footer_paragraph = footer.paragraphs[0]
    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_run = footer_paragraph.add_run("Gestion Projets IT - Côte d'Ivoire Terminal")
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = RGBColor(110, 110, 110)


def add_cover(doc: Document, title: str, subtitle: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if LOGO_PATH.exists():
        p.add_run().add_picture(str(LOGO_PATH), width=Inches(1.1))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(23)
    r.font.name = "Aptos Display"
    r.font.color.rgb = RGBColor(31, 45, 61)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(subtitle)
    r.italic = True
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(73, 95, 131)

    meta_lines = [
        "Application : Gestion Projets IT - Côte d'Ivoire Terminal",
        f"Version documentaire : {TODAY.strftime('%d/%m/%Y')}",
        "Principe : chaque écran décrit son rôle, son comportement attendu, les données à renseigner, les contrôles et les sorties.",
        "Branding : tout document généré nativement par l'application pour un projet doit porter le logo CIT.",
    ]
    for line in meta_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line)
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(86, 86, 86)
    doc.add_page_break()


def add_paragraph(doc: Document, text: str, *, bold_prefix: str | None = None) -> None:
    paragraph = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        run = paragraph.add_run(bold_prefix)
        run.bold = True
        paragraph.add_run(text[len(bold_prefix) :])
    else:
        paragraph.add_run(text)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.add_run(item)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], *, header_fill: str = "D9EAF7") -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_text(cell, header, bold=True)
        shade_cell(cell, header_fill)
    for row_data in rows:
        row = table.add_row().cells
        for idx, value in enumerate(row_data):
            set_cell_text(row[idx], value)
    doc.add_paragraph()


def add_screen_detail(doc: Document, screen: dict[str, str]) -> None:
    doc.add_heading(screen["name"], level=2)
    add_table(
        doc,
        ["Rubrique", "Contenu"],
        [
            ["Route / accès", screen["route"]],
            ["Acteurs", screen["roles"]],
            ["Objectif", screen["purpose"]],
            ["Comportement attendu", screen["behavior"]],
            ["Comment le renseigner", screen["fill"]],
            ["Sorties / effets", screen["outputs"]],
            ["Blocages / contrôles", screen["blockers"]],
        ],
        header_fill="E8EEF9",
    )


def collect_view_inventory() -> list[tuple[str, str, str, str]]:
    rows: list[tuple[str, str, str, str]] = []
    for path in sorted(VIEWS_ROOT.rglob("*.cshtml")):
        relative = path.relative_to(VIEWS_ROOT).as_posix()
        view_type, behavior, fill = VIEW_DESCRIPTIONS.get(
            relative,
            (
                "Vue",
                "Vue applicative à documenter plus finement si elle devient un écran métier principal.",
                "Se reporter à l'écran parent ou au contrôleur associé.",
            ),
        )
        rows.append((relative, view_type, behavior, fill))
    return rows


def add_view_inventory(doc: Document, title: str) -> None:
    doc.add_heading(title, level=1)
    add_paragraph(
        doc,
        "L'annexe suivante recense toutes les vues Razor présentes dans l'application afin de lier le document fonctionnel au périmètre réel du code. Les partiels et modales sont inclus pour expliciter le comportement attendu des sous-composants.",
    )
    grouped: dict[str, list[tuple[str, str, str, str]]] = defaultdict(list)
    for relative, view_type, behavior, fill in collect_view_inventory():
        grouped[relative.split("/", 1)[0]].append((relative, view_type, behavior, fill))

    for folder in sorted(grouped):
        doc.add_heading(folder, level=2)
        add_table(
            doc,
            ["Vue", "Type", "Comportement attendu", "Comment la renseigner / l'utiliser"],
            [list(item) for item in grouped[folder]],
            header_fill="EEF3F8",
        )


def build_workflow_doc() -> Path:
    doc = Document()
    apply_doc_defaults(doc, "Procédure DSI - Workflow par rôle Gestion Projets IT")
    add_header(doc.sections[0], "Procédure DSI - Workflow par rôle Gestion Projets IT")
    add_cover(
        doc,
        "Procédure DSI - Workflow par rôle",
        "Mode opératoire DSI centré sur les rôles, les validations, les écrans et les contrôles à appliquer dans l'application.",
    )

    doc.add_heading("1. Objet du document", level=1)
    add_paragraph(
        doc,
        "Ce document est rédigé comme une procédure d'exploitation DSI. Il décrit qui agit, sur quel écran, dans quel ordre, avec quels contrôles et avec quelle preuve attendue. Il s'appuie sur les écrans réellement présents, les validations effectivement câblées et les règles de blocage au 21/05/2026.",
    )
    add_paragraph(
        doc,
        "Convention importante : les documents produits nativement par l'application pour un projet (Word, PDF, Excel) doivent porter le logo CIT. Les pièces déposées manuellement conservent leur mise en forme d'origine. La procédure fait toujours la différence entre saisie native, pièce déposée et validation formelle.",
    )

    doc.add_heading("2. Matrice synthétique des rôles", level=1)
    add_table(doc, ["Rôle", "Mission", "Écrans principaux"], ROLE_MATRIX)

    doc.add_heading("3. Mode opératoire DSI - lecture rapide", level=1)
    add_table(
        doc,
        ["Étape", "Responsable principal", "Écran de travail", "Contrôle avant décision", "Preuve attendue"],
        [
            ["Demande initiale", "Demandeur -> DM -> DSI", "Create / Validations DM / Validations DSI", "Demande complète, pièces utiles, commentaire de décision si retour ou refus.", "Demande validée DSI et projet créé."],
            ["Analyse", "Chef de Projet -> DM -> DSI/RSIT", "Analyse / CharteProjet / ValidationsProjet", "Charte générée, charte signée complète, documents d'analyse présents.", "Validation DM puis validation DSI."],
            ["Planification", "Chef de Projet -> DM -> DSI/RSIT", "Planification", "Tâches saisies, livrables natifs générés, budget/gouvernance complétés.", "Passage en Exécution."],
            ["Exécution", "Chef de Projet", "Exécution / Charges", "Avancement à jour, charges renseignées, blocages et décisions tracés.", "Projet piloté et prêt pour recette."],
            ["UAT & MEP", "Chef de Projet + DSI", "UAT & MEP", "Tests, anomalies, plan de MEP, rollback et hypercare correctement renseignés.", "Projet prêt à clôture."],
            ["Clôture", "Chef de Projet -> Demandeur -> DM -> DSI", "Clôture + listes de validation", "Bilan final, transfert RUN, documentation et validations séquentielles.", "Projet clôturé."],
        ],
    )

    doc.add_heading("4. Workflow détaillé par rôle", level=1)
    for role_workflow in ROLE_WORKFLOWS:
        doc.add_heading(role_workflow["role"], level=2)
        add_paragraph(doc, role_workflow["mission"])
        add_table(
            doc,
            ["Écran", "Route", "Comportement attendu", "Action attendue", "Résultat"],
            [list(screen) for screen in role_workflow["screens"]],
            header_fill="E8EEF9",
        )

    doc.add_heading("5. Workflow détaillé par phase", level=1)
    for phase_name, intro, bullets in PHASE_GUIDES:
        doc.add_heading(phase_name, level=2)
        add_paragraph(doc, intro)
        add_bullets(doc, bullets)

    doc.add_heading("6. Focus : planification native", level=1)
    add_paragraph(
        doc,
        "La planification n'est plus limitée à l'upload de fichiers externes. Le Chef de Projet saisit maintenant directement dans l'application le planning interactif, les tâches, le WBS, la matrice RACI, le schéma de communication, le budget prévisionnel et le PV de kick-off. Le rôle de la DSI n'est plus de collecter des fichiers, mais de contrôler la cohérence du dossier natif avant décision.",
    )
    add_table(
        doc,
        ["Élément", "État actuel", "Mode d'usage"],
        [
            ["Planning détaillé", "Natif", "Saisie dans le Gantt puis génération Excel avec logo."],
            ["WBS", "Natif", "Dérivé des tâches du planning et généré automatiquement."],
            ["Matrice RACI", "Natif", "Saisie en grille directement dans la planification puis génération Excel."],
            ["Schéma de communication", "Natif", "Saisie des réunions, canaux et participants puis génération Excel."],
            ["Budget prévisionnel", "Natif", "Saisie des lignes budgétaires puis génération Excel."],
            ["PV de kick-off", "Natif", "Saisie des informations de réunion puis génération Excel."],
            ["Charte signée", "Semi-natif", "La charte est générée par l'application, mais la version signée reste un dépôt externe."],
        ],
    )

    doc.add_heading("7. Checklist de renseignement par phase", level=1)
    add_table(
        doc,
        ["Phase", "Ce que le CP renseigne directement", "Ce que le valideur contrôle", "Sortie de phase"],
        [
            ["Analyse", "Contexte, objectifs, cahier d'analyse, note de cadrage, charte, signatures.", "Le valideur contrôle la charte signée, le contenu du cadrage et la cohérence métier.", "Passage en Planification."],
            ["Planification", "Gantt, WBS, RACI, communication, budget, PV de kick-off.", "Le valideur contrôle la cohérence du planning, des jalons, des responsabilités et du budget.", "Passage en Exécution."],
            ["Exécution", "Avancement, décisions, blocages, charges, risques, changements.", "Le pilotage contrôle la fiabilité des données et les écarts.", "Ouverture UAT / MEP."],
            ["UAT & MEP", "Tests, anomalies, MEP, rollback, incidents, hypercare.", "Le valideur contrôle la maîtrise du risque de mise en production.", "Ouverture Clôture."],
            ["Clôture", "Bilan final, transfert RUN, documentation, leçons apprises.", "Les valideurs contrôlent le bilan, le transfert et la fin opérationnelle réelle.", "Projet clôturé."],
        ],
    )

    add_view_inventory(doc, "8. Annexe - inventaire des vues et composants")
    return save_with_fallback(doc, WORKFLOW_DOC)


def build_complete_doc() -> Path:
    doc = Document()
    apply_doc_defaults(doc, "Manuel utilisateur - Gestion Projets IT")
    add_header(doc.sections[0], "Manuel utilisateur - Gestion Projets IT")
    add_cover(
        doc,
        "Manuel utilisateur de l'application",
        "Guide d'utilisation écran par écran de Gestion Projets IT, rédigé pour les utilisateurs finaux et les acteurs métier.",
    )

    doc.add_heading("1. Comment utiliser ce manuel", level=1)
    add_paragraph(
        doc,
        "Ce document se lit comme un manuel utilisateur. Pour chaque écran principal, il explique à quoi l'écran sert, ce qu'il affiche, ce que l'utilisateur doit y saisir, ce qui déclenche un blocage et le résultat attendu après action.",
    )
    add_bullets(
        doc,
        [
            "Si vous êtes demandeur, commencez par les chapitres Demandes puis Clôture demandeur.",
            "Si vous êtes Chef de Projet, concentrez-vous sur Détail projet, Charte, Planification, Exécution, Charges, UAT & MEP et Clôture.",
            "Si vous êtes valideur métier, utilisez les écrans de validation correspondants et les listes d'attente.",
            "Si vous êtes AdminIT, utilisez en plus les chapitres Administration et Autorisations.",
        ],
    )

    doc.add_heading("2. Positionnement de l'application", level=1)
    add_paragraph(
        doc,
        "Gestion Projets IT centralise le cycle complet d'un projet IT : demande initiale, analyse, charte, planification, exécution, UAT / mise en production, clôture, portefeuille, analytics, administration et sécurité.",
    )
    add_bullets(
        doc,
        [
            "Chaque rôle voit un menu différent, piloté par la matrice des autorisations et par les règles métier.",
            "Les validations structurantes se font dans l'application et sont tracées.",
            "Les documents générés nativement pour le projet portent le logo CIT.",
            "L'objectif produit est d'éviter la double saisie et de rendre le dossier projet le plus natif possible.",
        ],
    )

    doc.add_heading("3. Principes d'usage", level=1)
    add_table(
        doc,
        ["Principe", "Application pratique"],
        [
            ["Natif", "La donnée est saisie directement dans un écran et conservée en base ; l'application peut ensuite générer le livrable officiel."],
            ["Semi-natif", "Une partie du flux est saisie dans l'application, mais il reste un dépôt de pièce externe ou signée."],
            ["Manuel", "Le document est produit hors application puis déposé comme livrable."],
            ["Rôle de vérité", "Le dernier état validé dans l'écran métier fait foi pour le workflow et les blocages."],
        ],
    )

    doc.add_heading("4. Démarrage rapide par rôle", level=1)
    add_table(
        doc,
        ["Rôle", "Par quoi commencer", "Ce que vous faites le plus souvent"],
        [
            ["Demandeur", "Mes Demandes / Nouvelle Demande", "Créer une demande, suivre les retours, valider la clôture demandeur."],
            ["Directeur Métier", "Validations DM / Validations Projet", "Valider demandes, charte, planification et clôture sponsor."],
            ["Chef de Projet", "Mes projets / Détail projet", "Renseigner toutes les phases et produire les livrables natifs."],
            ["DSI", "Dashboard / Validations DSI / Validations Projet", "Arbitrer, valider et piloter le portefeuille."],
            ["RSIT", "Portefeuille / projets / validations déléguées", "Consulter, aider à l'instruction et valider si délégation active."],
            ["AdminIT", "Utilisateurs / Autorisations / Paramètres", "Gérer les comptes, les droits et le paramétrage."],
        ],
    )

    doc.add_heading("5. Architecture fonctionnelle", level=1)
    add_table(
        doc,
        ["Module", "But", "Écrans principaux"],
        [
            ["Accès & shell", "Connexion, navigation, profil, notifications.", "Login, Profil, Notifications, Menu latéral, Aide."],
            ["Demandes", "Créer et valider le besoin initial.", "Mes Demandes, Nouvelle Demande, Validations DM, Validations DSI, Historiques."],
            ["Projets", "Piloter toutes les phases du projet.", "Détail projet, Charte, Fiche projet, Charges, Portefeuille, Validations projet."],
            ["Administration", "Gérer utilisateurs, rôles, paramètres et droits.", "Utilisateurs, Rôles, Directions, Services, Délégations, Autorisations."],
            ["Reporting", "Observer le portefeuille et les alertes.", "Dashboard, Portefeuille, Analytics, Historiques."],
        ],
    )

    doc.add_heading("6. Écrans détaillés", level=1)
    for screen in DETAILED_SCREENS:
        add_screen_detail(doc, screen)

    doc.add_heading("7. Règles de blocage et validations", level=1)
    add_table(
        doc,
        ["Sujet", "Règle actuelle"],
        [
            ["Validation de la charte", "DM puis DSI/RSIT délégué. La charte signée doit être complète (document + signature sponsor + signature chef de projet)."],
            ["Passage Analyse -> Planification", "Charte projet générée, charte signée complète, validation DM et validation DSI."],
            ["Passage Planification -> Exécution", "Planning interactif enregistré, livrables natifs générés, validation DM puis validation DSI/RSIT délégué."],
            ["RSIT", "Peut valider à la place de la DSI uniquement si une délégation active existe."],
            ["AdminIT", "Conserve un accès complet et n'est pas restreint par la matrice pour les routes critiques."],
        ],
    )

    doc.add_heading("8. Livrables par phase", level=1)
    add_table(
        doc,
        ["Phase", "Livrables attendus", "Mode actuel"],
        [
            ["Analyse", "Cahier d'analyse technique, Note de cadrage, Charte projet, Charte projet signée.", "Analyse : uploads + génération PDF charte + dépôt signé."],
            ["Planification", "Planning détaillé, WBS, Matrice RACI, Schéma de communication, Budget prévisionnel, PV de kick-off.", "Natif avec génération Excel brandée."],
            ["Exécution", "Comptes-rendus, supports de suivi, états d'avancement, pièces opérationnelles.", "Principalement natif sur la donnée, upload pour pièces externes."],
            ["UAT & MEP", "Campagnes, cas de test, plan de MEP, rollback, bilan UAT, incidents.", "Mix de saisie native et dépôts complémentaires."],
            ["Clôture", "Bilan final, transfert RUN, pièces de clôture.", "Natif pour le bilan, dépôts complémentaires si nécessaire."],
        ],
    )

    doc.add_heading("9. Règles documentaires", level=1)
    add_bullets(
        doc,
        [
            "Les documents générés par le projet (Word, PDF, Excel) doivent inclure le logo CIT.",
            "Le branding s'applique aux générations natives : charte, fiche projet, rapports, planning, WBS, RACI, communication, budget, PV de kick-off.",
            "Les fichiers déposés manuellement restent inchangés ; ils ne sont pas rebrandés par l'application.",
        ],
    )

    doc.add_heading("10. Sécurité et autorisations", level=1)
    add_bullets(
        doc,
        [
            "L'accès repose sur l'authentification locale autorisée ou Azure AD selon la configuration active.",
            "La matrice des autorisations pilote maintenant à la fois le menu visible et le blocage backend par route.",
            "Les validations métier sont renforcées par des contrôles serveurs : les boutons masqués ou grisés ne suffisent pas à eux seuls, le backend vérifie aussi les prérequis.",
            "La politique mot de passe est harmonisée sur les écrans de création, modification et réinitialisation.",
        ],
    )

    add_view_inventory(doc, "11. Annexe - inventaire complet des vues Razor")
    return save_with_fallback(doc, DOCUMENTATION_DOC)


def main() -> None:
    workflow_path = build_workflow_doc()
    documentation_path = build_complete_doc()
    print(f"Updated: {workflow_path}")
    print(f"Updated: {documentation_path}")


if __name__ == "__main__":
    main()
